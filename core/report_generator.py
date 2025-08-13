"""
Comprehensive Report Generator with multiple export formats and interactive visualizations
"""

import logging
import os
import json
import csv
import time
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple, Union
from datetime import datetime, timedelta
from collections import defaultdict, Counter
import base64
import io

import numpy as np
import pandas as pd
from scipy import stats
from scipy.stats import pearsonr

try:
    import plotly.graph_objects as go
    import plotly.express as px
    import plotly.figure_factory as ff
    from plotly.subplots import make_subplots
    import plotly.io as pio
    PLOTLY_AVAILABLE = True
except ImportError:
    PLOTLY_AVAILABLE = False

try:
    import weasyprint
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Internal imports
from utils.translation import EmotionTranslator


class ComprehensiveReportGenerator:
    """
    Comprehensive report generator with multiple formats and interactive visualizations
    for interrogation analysis
    """
    
    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.logger = logging.getLogger(__name__)
        
        # Report configuration
        report_config = config.get('analysis', {}).get('report', {})
        self.output_formats = report_config.get('formats', ['csv', 'json', 'html'])
        self.include_interactive = report_config.get('interactive', True)
        self.include_statistics = report_config.get('statistics', True)
        
        # Storage paths
        storage_config = config.get('storage', {})
        self.reports_dir = Path(storage_config.get('reports_dir', 'storage/reports'))
        self.assets_dir = self.reports_dir / 'assets'
        
        # Create directories
        self.reports_dir.mkdir(parents=True, exist_ok=True)
        self.assets_dir.mkdir(parents=True, exist_ok=True)
        
        # Visualization settings
        viz_config = report_config.get('visualization', {})
        self.color_scheme = viz_config.get('color_scheme', 'viridis')
        self.figure_width = viz_config.get('figure_width', 1200)
        self.figure_height = viz_config.get('figure_height', 600)
        
        # Initialize components
        self.emotion_translator = EmotionTranslator()
        
        # Set plotly theme
        if PLOTLY_AVAILABLE:
            pio.templates.default = "plotly_white"
        
        self.logger.info("ComprehensiveReportGenerator initialized")
    
    def generate_full_report(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Generate complete analysis report with all visualizations and exports
        
        Args:
            data: Complete analysis data including:
                - video_emotions: Video emotion analysis results
                - speech_emotions: Speech emotion analysis results
                - transcript: Audio transcript data
                - critical_moments: Detected critical moments
                - openai_insights: GPT-4 analysis results
                
        Returns:
            Dict with all generated reports and file paths
        """
        try:
            self.logger.info("Starting comprehensive report generation")
            
            # Create timestamp for this report
            timestamp = int(time.time())
            report_id = f"interrogation_analysis_{timestamp}"
            
            # Prepare data for analysis
            processed_data = self._prepare_data(data)
            
            # Generate all components
            results = {
                'report_id': report_id,
                'timestamp': timestamp,
                'generated_at': datetime.now().isoformat(),
                'formats': {},
                'visualizations': {},
                'statistics': {},
                'files': []
            }
            
            # 1. Generate CSV reports
            if 'csv' in self.output_formats:
                csv_reports = self.generate_csv_reports(processed_data, report_id)
                results['formats']['csv'] = csv_reports
                results['files'].extend(csv_reports.values())
            
            # 2. Generate interactive visualizations
            if PLOTLY_AVAILABLE and self.include_interactive:
                visualizations = self._generate_all_visualizations(processed_data)
                results['visualizations'] = visualizations
            
            # 3. Generate HTML report
            if 'html' in self.output_formats:
                html_path = self.generate_html_report(processed_data, results['visualizations'], report_id)
                results['formats']['html'] = html_path
                results['files'].append(html_path)
            
            # 4. Generate JSON report
            if 'json' in self.output_formats:
                json_path = self.generate_json_report(processed_data, report_id)
                results['formats']['json'] = json_path
                results['files'].append(json_path)
            
            # 5. Generate PDF report
            if 'pdf' in self.output_formats:
                pdf_path = self.generate_pdf_report(processed_data, results['visualizations'], report_id)
                if pdf_path:
                    results['formats']['pdf'] = pdf_path
                    results['files'].append(pdf_path)
            
            # 6. Generate DOCX report
            if 'docx' in self.output_formats:
                docx_path = self.generate_docx_report(processed_data, report_id)
                if docx_path:
                    results['formats']['docx'] = docx_path
                    results['files'].append(docx_path)
            
            # 7. Generate statistics
            if self.include_statistics:
                statistics = self._generate_comprehensive_statistics(processed_data)
                results['statistics'] = statistics
            
            self.logger.info(f"Report generation completed: {len(results['files'])} files created")
            return results
            
        except Exception as e:
            self.logger.error(f"Report generation failed: {e}")
            raise
    
    def _prepare_data(self, raw_data: Dict[str, Any]) -> Dict[str, Any]:
        """Prepare and normalize data for report generation"""
        
        processed = {
            'metadata': {
                'duration': 0,
                'video_frames': 0,
                'audio_segments': 0,
                'processing_time': time.time()
            },
            'video_emotions': [],
            'speech_emotions': [],
            'transcript': [],
            'critical_moments': [],
            'openai_insights': {},
            'correlations': {},
            'statistics': {}
        }
        
        try:
            # Process video emotions
            if 'video_emotions' in raw_data:
                processed['video_emotions'] = self._normalize_emotion_data(
                    raw_data['video_emotions'], 'video'
                )
                processed['metadata']['video_frames'] = len(processed['video_emotions'])
            
            # Process speech emotions
            if 'speech_emotions' in raw_data:
                processed['speech_emotions'] = self._normalize_emotion_data(
                    raw_data['speech_emotions'], 'speech'
                )
                processed['metadata']['audio_segments'] = len(processed['speech_emotions'])
            
            # Process transcript
            if 'transcript' in raw_data:
                processed['transcript'] = raw_data['transcript']
            
            # Process critical moments
            if 'critical_moments' in raw_data:
                processed['critical_moments'] = raw_data['critical_moments']
            
            # Process OpenAI insights
            if 'openai_insights' in raw_data:
                processed['openai_insights'] = raw_data['openai_insights']
            
            # Calculate duration
            all_timestamps = []
            if processed['video_emotions']:
                all_timestamps.extend([item.get('timestamp', 0) for item in processed['video_emotions']])
            if processed['speech_emotions']:
                all_timestamps.extend([item.get('timestamp', 0) for item in processed['speech_emotions']])
            
            if all_timestamps:
                processed['metadata']['duration'] = max(all_timestamps)
            
            # Calculate correlations
            processed['correlations'] = self._calculate_correlations(processed)
            
            # Generate basic statistics
            processed['statistics'] = self._calculate_basic_statistics(processed)
            
        except Exception as e:
            self.logger.warning(f"Data preparation warning: {e}")
        
        return processed
    
    def _normalize_emotion_data(self, emotion_data: List[Dict], source: str) -> List[Dict]:
        """Normalize emotion data to consistent format"""
        normalized = []
        
        for item in emotion_data:
            if not isinstance(item, dict):
                continue
            
            normalized_item = {
                'timestamp': item.get('timestamp', 0),
                'emotion': item.get('emotion', 'нейтральность'),
                'emotion_en': self.emotion_translator.reverse_translate(item.get('emotion', 'нейтральность')),
                'confidence': item.get('confidence', 0.0),
                'source': source
            }
            
            # Add source-specific fields
            if source == 'video':
                normalized_item.update({
                    'frame_number': item.get('frame_number', 0),
                    'face_bbox': item.get('face_bbox'),
                    'method': item.get('method', 'unknown')
                })
            elif source == 'speech':
                normalized_item.update({
                    'segment_start': item.get('segment_start', 0),
                    'segment_end': item.get('segment_end', 0),
                    'text': item.get('text', ''),
                    'speaker': item.get('speaker', 'unknown')
                })
            
            normalized.append(normalized_item)
        
        return sorted(normalized, key=lambda x: x['timestamp'])
    
    def generate_csv_reports(self, data: Dict[str, Any], report_id: str) -> Dict[str, str]:
        """Generate multiple CSV reports"""
        
        csv_reports = {}
        
        try:
            # 1. Emotions timeline
            timeline_path = self._create_emotions_timeline(data, report_id)
            csv_reports['emotions_timeline'] = timeline_path
            
            # 2. Transitions report
            transitions_path = self._create_transitions_report(data, report_id)
            csv_reports['transitions'] = transitions_path
            
            # 3. Speech emotions report
            speech_path = self._create_speech_report(data, report_id)
            csv_reports['speech_emotions'] = speech_path
            
            # 4. Critical moments
            critical_path = self._create_critical_moments_report(data, report_id)
            csv_reports['critical_moments'] = critical_path
            
            # 5. Correlation matrix
            correlation_path = self._create_correlation_matrix(data, report_id)
            csv_reports['correlation_matrix'] = correlation_path
            
            # 6. Summary statistics
            stats_path = self._create_statistics_report(data, report_id)
            csv_reports['summary_statistics'] = stats_path
            
        except Exception as e:
            self.logger.error(f"CSV reports generation failed: {e}")
        
        return csv_reports
    
    def _create_emotions_timeline(self, data: Dict[str, Any], report_id: str) -> str:
        """Create emotions timeline CSV"""
        
        output_path = str(self.reports_dir / f"{report_id}_emotions_timeline.csv")
        
        # Combine video and speech emotions
        timeline_data = []
        
        for item in data.get('video_emotions', []):
            timeline_data.append({
                'timestamp': item['timestamp'],
                'source': 'video',
                'emotion_ru': item['emotion'],
                'emotion_en': item['emotion_en'],
                'confidence': item['confidence'],
                'frame_number': item.get('frame_number', ''),
                'method': item.get('method', ''),
                'speaker': '',
                'text': ''
            })
        
        for item in data.get('speech_emotions', []):
            timeline_data.append({
                'timestamp': item['timestamp'],
                'source': 'speech',
                'emotion_ru': item['emotion'],
                'emotion_en': item['emotion_en'],
                'confidence': item['confidence'],
                'frame_number': '',
                'method': '',
                'speaker': item.get('speaker', ''),
                'text': item.get('text', '')
            })
        
        # Sort by timestamp
        timeline_data.sort(key=lambda x: x['timestamp'])
        
        # Write to CSV
        if timeline_data:
            df = pd.DataFrame(timeline_data)
            df.to_csv(output_path, index=False, encoding='utf-8')
        
        return output_path
    
    def _generate_summary_statistics(self, data: Dict[str, Any]) -> str:
        """Generate HTML for summary statistics"""
        metadata = data.get('metadata', {})
        
        # Calculate emotion distribution
        video_emotions = data.get('video_emotions', [])
        speech_emotions = data.get('speech_emotions', [])
        
        video_emotion_dist = Counter([item['emotion'] for item in video_emotions])
        speech_emotion_dist = Counter([item['emotion'] for item in speech_emotions])
        
        # Most common emotions
        most_common_video = video_emotion_dist.most_common(1)[0] if video_emotion_dist else ('нет данных', 0)
        most_common_speech = speech_emotion_dist.most_common(1)[0] if speech_emotion_dist else ('нет данных', 0)
        
        # Calculate average confidence
        video_confidences = [item['confidence'] for item in video_emotions]
        speech_confidences = [item['confidence'] for item in speech_emotions]
        
        avg_video_conf = np.mean(video_confidences) if video_confidences else 0
        avg_speech_conf = np.mean(speech_confidences) if speech_confidences else 0
        
        # Generate HTML cards
        stats_html = f"""
        <div class="stat-card">
            <div class="stat-value">{metadata.get('duration', 0):.1f}с</div>
            <div>Общая продолжительность</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{len(video_emotions)}</div>
            <div>Анализов лица</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{len(speech_emotions)}</div>
            <div>Сегментов речи</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{most_common_video[0]}</div>
            <div>Частая эмоция (лицо)</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{most_common_speech[0]}</div>
            <div>Частая эмоция (речь)</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{avg_video_conf:.2f}</div>
            <div>Средняя уверенность (лицо)</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{avg_speech_conf:.2f}</div>
            <div>Средняя уверенность (речь)</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{len(data.get('critical_moments', []))}</div>
            <div>Критических моментов</div>
        </div>
        """
        
        return stats_html
    
    def _generate_critical_moments_html(self, data: Dict[str, Any]) -> str:
        """Generate HTML for critical moments"""
        critical_moments = data.get('critical_moments', [])
        
        if not critical_moments:
            return '<p>Критические моменты не обнаружены или не анализировались.</p>'
        
        moments_html = ''
        for moment in sorted(critical_moments, key=lambda x: x.get('timestamp', 0)):
            timestamp = moment.get('timestamp', 0)
            time_str = f"{int(timestamp//60)}:{int(timestamp%60):02d}"
            importance = moment.get('importance', 5)
            description = moment.get('description', 'Описание отсутствует')
            context = moment.get('context', '')
            emotions = moment.get('emotions', [])
            
            emotion_tags = ''.join([f'<span class="emotion-tag">{emotion}</span>' for emotion in emotions])
            
            moments_html += f"""
            <div class="critical-moment">
                <h4>⚠️ {time_str} - Важность: {importance}/10</h4>
                <p><strong>Описание:</strong> {description}</p>
                {f'<p><strong>Контекст:</strong> {context}</p>' if context else ''}
                {f'<p><strong>Связанные эмоции:</strong> {emotion_tags}</p>' if emotions else ''}
            </div>
            """
        
        return moments_html
    
    def _generate_transcript_html(self, data: Dict[str, Any]) -> str:
        """Generate HTML for transcript with emotion analysis"""
        speech_emotions = data.get('speech_emotions', [])
        transcript = data.get('transcript', [])
        
        if not speech_emotions and not transcript:
            return '<p>Транскрипция недоступна.</p>'
        
        transcript_html = ''
        
        # Use speech emotions if available, otherwise fallback to transcript
        segments = speech_emotions if speech_emotions else transcript
        
        for segment in segments:
            timestamp = segment.get('timestamp', 0)
            time_str = f"{int(timestamp//60)}:{int(timestamp%60):02d}"
            text = segment.get('text', 'Текст недоступен')
            speaker = segment.get('speaker', 'Неизвестный')
            emotion = segment.get('emotion', 'неопределено')
            confidence = segment.get('confidence', 0)
            
            # Get emotion color
            emotion_color = self.emotion_translator.get_emotion_color(emotion)
            
            transcript_html += f"""
            <div class="transcript-segment">
                <div class="timestamp">{time_str} - {speaker}</div>
                <p><strong>{text}</strong></p>
                <div style="margin-top: 10px;">
                    <span class="emotion-tag" style="background-color: {emotion_color}">
                        {emotion} ({confidence:.2f})
                    </span>
                </div>
            </div>
            """
        
        return transcript_html
    
    def _generate_insights_html(self, data: Dict[str, Any]) -> str:
        """Generate HTML for psychological insights"""
        insights = data.get('openai_insights', {})
        
        if not insights:
            return '<p>Психологические инсайты недоступны. Требуется интеграция с OpenAI.</p>'
        
        insights_html = ''
        
        # General analysis
        if 'general_analysis' in insights:
            insights_html += f"""
            <h3>🎯 Общий анализ</h3>
            <p>{insights['general_analysis']}</p>
            """
        
        # Emotional patterns
        if 'emotional_patterns' in insights:
            insights_html += f"""
            <h3>📈 Эмоциональные паттерны</h3>
            <p>{insights['emotional_patterns']}</p>
            """
        
        # Behavioral indicators
        if 'behavioral_indicators' in insights:
            insights_html += f"""
            <h3>🎭 Поведенческие индикаторы</h3>
            <p>{insights['behavioral_indicators']}</p>
            """
        
        # Recommendations
        if 'recommendations' in insights:
            insights_html += f"""
            <h3>💡 Рекомендации</h3>
            <p>{insights['recommendations']}</p>
            """
        
        # Risk assessment
        if 'risk_assessment' in insights:
            insights_html += f"""
            <h3>⚡ Оценка рисков</h3>
            <p>{insights['risk_assessment']}</p>
            """
        
        return insights_html if insights_html else '<p>Инсайты обрабатываются...</p>'
    
    def generate_json_report(self, data: Dict[str, Any], report_id: str) -> str:
        """Generate comprehensive JSON report"""
        output_path = str(self.reports_dir / f"{report_id}_report.json")
        
        json_data = {
            'metadata': {
                'report_id': report_id,
                'generated_at': datetime.now().isoformat(),
                'version': 'ДОПРОС MVP 2.0',
                'duration': data.get('metadata', {}).get('duration', 0),
                'video_frames': data.get('metadata', {}).get('video_frames', 0),
                'audio_segments': data.get('metadata', {}).get('audio_segments', 0)
            },
            'video_emotions': data.get('video_emotions', []),
            'speech_emotions': data.get('speech_emotions', []),
            'critical_moments': data.get('critical_moments', []),
            'transcript': data.get('transcript', []),
            'correlations': data.get('correlations', {}),
            'statistics': data.get('statistics', {}),
            'openai_insights': data.get('openai_insights', {})
        }
        
        # Add computed statistics
        video_emotions = data.get('video_emotions', [])
        speech_emotions = data.get('speech_emotions', [])
        
        if video_emotions:
            json_data['computed_statistics'] = {
                'video_emotion_distribution': dict(Counter([item['emotion'] for item in video_emotions])),
                'video_avg_confidence': float(np.mean([item['confidence'] for item in video_emotions])),
                'video_emotion_changes': self._count_emotion_changes(video_emotions)
            }
        
        if speech_emotions:
            json_data['computed_statistics'].update({
                'speech_emotion_distribution': dict(Counter([item['emotion'] for item in speech_emotions])),
                'speech_avg_confidence': float(np.mean([item['confidence'] for item in speech_emotions])),
                'speech_emotion_changes': self._count_emotion_changes(speech_emotions)
            })
        
        # Write JSON file
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, ensure_ascii=False, indent=2)
        
        return output_path
    
    def _create_transitions_report(self, data: Dict[str, Any], report_id: str) -> str:
        """Create emotion transitions CSV"""
        
        output_path = str(self.reports_dir / f"{report_id}_transitions.csv")
        
        transitions_data = []
        
        # Analyze video emotion transitions
        video_emotions = data.get('video_emotions', [])
        for i in range(1, len(video_emotions)):
            prev_item = video_emotions[i-1]
            curr_item = video_emotions[i]
            
            if prev_item['emotion'] != curr_item['emotion']:
                transitions_data.append({
                    'timestamp': curr_item['timestamp'],
                    'source': 'video',
                    'from_emotion': prev_item['emotion'],
                    'to_emotion': curr_item['emotion'],
                    'from_confidence': prev_item['confidence'],
                    'to_confidence': curr_item['confidence'],
                    'duration': curr_item['timestamp'] - prev_item['timestamp']
                })
        
        # Analyze speech emotion transitions
        speech_emotions = data.get('speech_emotions', [])
        for i in range(1, len(speech_emotions)):
            prev_item = speech_emotions[i-1]
            curr_item = speech_emotions[i]
            
            if prev_item['emotion'] != curr_item['emotion']:
                transitions_data.append({
                    'timestamp': curr_item['timestamp'],
                    'source': 'speech',
                    'from_emotion': prev_item['emotion'],
                    'to_emotion': curr_item['emotion'],
                    'from_confidence': prev_item['confidence'],
                    'to_confidence': curr_item['confidence'],
                    'duration': curr_item['timestamp'] - prev_item['timestamp']
                })
        
        # Write to CSV
        if transitions_data:
            df = pd.DataFrame(transitions_data)
            df.to_csv(output_path, index=False, encoding='utf-8')
        
        return output_path
    
    def _generate_summary_statistics(self, data: Dict[str, Any]) -> str:
        """Generate HTML for summary statistics"""
        metadata = data.get('metadata', {})
        
        # Calculate emotion distribution
        video_emotions = data.get('video_emotions', [])
        speech_emotions = data.get('speech_emotions', [])
        
        video_emotion_dist = Counter([item['emotion'] for item in video_emotions])
        speech_emotion_dist = Counter([item['emotion'] for item in speech_emotions])
        
        # Most common emotions
        most_common_video = video_emotion_dist.most_common(1)[0] if video_emotion_dist else ('нет данных', 0)
        most_common_speech = speech_emotion_dist.most_common(1)[0] if speech_emotion_dist else ('нет данных', 0)
        
        # Calculate average confidence
        video_confidences = [item['confidence'] for item in video_emotions]
        speech_confidences = [item['confidence'] for item in speech_emotions]
        
        avg_video_conf = np.mean(video_confidences) if video_confidences else 0
        avg_speech_conf = np.mean(speech_confidences) if speech_confidences else 0
        
        # Generate HTML cards
        stats_html = f"""
        <div class="stat-card">
            <div class="stat-value">{metadata.get('duration', 0):.1f}с</div>
            <div>Общая продолжительность</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{len(video_emotions)}</div>
            <div>Анализов лица</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{len(speech_emotions)}</div>
            <div>Сегментов речи</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{most_common_video[0]}</div>
            <div>Частая эмоция (лицо)</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{most_common_speech[0]}</div>
            <div>Частая эмоция (речь)</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{avg_video_conf:.2f}</div>
            <div>Средняя уверенность (лицо)</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{avg_speech_conf:.2f}</div>
            <div>Средняя уверенность (речь)</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{len(data.get('critical_moments', []))}</div>
            <div>Критических моментов</div>
        </div>
        """
        
        return stats_html
    
    def _generate_critical_moments_html(self, data: Dict[str, Any]) -> str:
        """Generate HTML for critical moments"""
        critical_moments = data.get('critical_moments', [])
        
        if not critical_moments:
            return '<p>Критические моменты не обнаружены или не анализировались.</p>'
        
        moments_html = ''
        for moment in sorted(critical_moments, key=lambda x: x.get('timestamp', 0)):
            timestamp = moment.get('timestamp', 0)
            time_str = f"{int(timestamp//60)}:{int(timestamp%60):02d}"
            importance = moment.get('importance', 5)
            description = moment.get('description', 'Описание отсутствует')
            context = moment.get('context', '')
            emotions = moment.get('emotions', [])
            
            emotion_tags = ''.join([f'<span class="emotion-tag">{emotion}</span>' for emotion in emotions])
            
            moments_html += f"""
            <div class="critical-moment">
                <h4>⚠️ {time_str} - Важность: {importance}/10</h4>
                <p><strong>Описание:</strong> {description}</p>
                {f'<p><strong>Контекст:</strong> {context}</p>' if context else ''}
                {f'<p><strong>Связанные эмоции:</strong> {emotion_tags}</p>' if emotions else ''}
            </div>
            """
        
        return moments_html
    
    def _generate_transcript_html(self, data: Dict[str, Any]) -> str:
        """Generate HTML for transcript with emotion analysis"""
        speech_emotions = data.get('speech_emotions', [])
        transcript = data.get('transcript', [])
        
        if not speech_emotions and not transcript:
            return '<p>Транскрипция недоступна.</p>'
        
        transcript_html = ''
        
        # Use speech emotions if available, otherwise fallback to transcript
        segments = speech_emotions if speech_emotions else transcript
        
        for segment in segments:
            timestamp = segment.get('timestamp', 0)
            time_str = f"{int(timestamp//60)}:{int(timestamp%60):02d}"
            text = segment.get('text', 'Текст недоступен')
            speaker = segment.get('speaker', 'Неизвестный')
            emotion = segment.get('emotion', 'неопределено')
            confidence = segment.get('confidence', 0)
            
            # Get emotion color
            emotion_color = self.emotion_translator.get_emotion_color(emotion)
            
            transcript_html += f"""
            <div class="transcript-segment">
                <div class="timestamp">{time_str} - {speaker}</div>
                <p><strong>{text}</strong></p>
                <div style="margin-top: 10px;">
                    <span class="emotion-tag" style="background-color: {emotion_color}">
                        {emotion} ({confidence:.2f})
                    </span>
                </div>
            </div>
            """
        
        return transcript_html
    
    def _generate_insights_html(self, data: Dict[str, Any]) -> str:
        """Generate HTML for psychological insights"""
        insights = data.get('openai_insights', {})
        
        if not insights:
            return '<p>Психологические инсайты недоступны. Требуется интеграция с OpenAI.</p>'
        
        insights_html = ''
        
        # General analysis
        if 'general_analysis' in insights:
            insights_html += f"""
            <h3>🎯 Общий анализ</h3>
            <p>{insights['general_analysis']}</p>
            """
        
        # Emotional patterns
        if 'emotional_patterns' in insights:
            insights_html += f"""
            <h3>📈 Эмоциональные паттерны</h3>
            <p>{insights['emotional_patterns']}</p>
            """
        
        # Behavioral indicators
        if 'behavioral_indicators' in insights:
            insights_html += f"""
            <h3>🎭 Поведенческие индикаторы</h3>
            <p>{insights['behavioral_indicators']}</p>
            """
        
        # Recommendations
        if 'recommendations' in insights:
            insights_html += f"""
            <h3>💡 Рекомендации</h3>
            <p>{insights['recommendations']}</p>
            """
        
        # Risk assessment
        if 'risk_assessment' in insights:
            insights_html += f"""
            <h3>⚡ Оценка рисков</h3>
            <p>{insights['risk_assessment']}</p>
            """
        
        return insights_html if insights_html else '<p>Инсайты обрабатываются...</p>'
    
    def generate_json_report(self, data: Dict[str, Any], report_id: str) -> str:
        """Generate comprehensive JSON report"""
        output_path = str(self.reports_dir / f"{report_id}_report.json")
        
        json_data = {
            'metadata': {
                'report_id': report_id,
                'generated_at': datetime.now().isoformat(),
                'version': 'ДОПРОС MVP 2.0',
                'duration': data.get('metadata', {}).get('duration', 0),
                'video_frames': data.get('metadata', {}).get('video_frames', 0),
                'audio_segments': data.get('metadata', {}).get('audio_segments', 0)
            },
            'video_emotions': data.get('video_emotions', []),
            'speech_emotions': data.get('speech_emotions', []),
            'critical_moments': data.get('critical_moments', []),
            'transcript': data.get('transcript', []),
            'correlations': data.get('correlations', {}),
            'statistics': data.get('statistics', {}),
            'openai_insights': data.get('openai_insights', {})
        }
        
        # Add computed statistics
        video_emotions = data.get('video_emotions', [])
        speech_emotions = data.get('speech_emotions', [])
        
        if video_emotions:
            json_data['computed_statistics'] = {
                'video_emotion_distribution': dict(Counter([item['emotion'] for item in video_emotions])),
                'video_avg_confidence': float(np.mean([item['confidence'] for item in video_emotions])),
                'video_emotion_changes': self._count_emotion_changes(video_emotions)
            }
        
        if speech_emotions:
            json_data['computed_statistics'].update({
                'speech_emotion_distribution': dict(Counter([item['emotion'] for item in speech_emotions])),
                'speech_avg_confidence': float(np.mean([item['confidence'] for item in speech_emotions])),
                'speech_emotion_changes': self._count_emotion_changes(speech_emotions)
            })
        
        # Write JSON file
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, ensure_ascii=False, indent=2)
        
        return output_path
    
    def _create_speech_report(self, data: Dict[str, Any], report_id: str) -> str:
        """Create detailed speech emotions CSV"""
        
        output_path = str(self.reports_dir / f"{report_id}_speech_emotions.csv")
        
        speech_data = []
        for item in data.get('speech_emotions', []):
            speech_data.append({
                'timestamp': item['timestamp'],
                'segment_start': item.get('segment_start', 0),
                'segment_end': item.get('segment_end', 0),
                'duration': item.get('segment_end', 0) - item.get('segment_start', 0),
                'emotion_ru': item['emotion'],
                'emotion_en': item['emotion_en'],
                'confidence': item['confidence'],
                'speaker': item.get('speaker', ''),
                'text': item.get('text', ''),
                'text_length': len(item.get('text', '')),
                'words_count': len(item.get('text', '').split()) if item.get('text') else 0
            })
        
        # Write to CSV
        if speech_data:
            df = pd.DataFrame(speech_data)
            df.to_csv(output_path, index=False, encoding='utf-8')
        
        return output_path
    
    def _generate_summary_statistics(self, data: Dict[str, Any]) -> str:
        """Generate HTML for summary statistics"""
        metadata = data.get('metadata', {})
        
        # Calculate emotion distribution
        video_emotions = data.get('video_emotions', [])
        speech_emotions = data.get('speech_emotions', [])
        
        video_emotion_dist = Counter([item['emotion'] for item in video_emotions])
        speech_emotion_dist = Counter([item['emotion'] for item in speech_emotions])
        
        # Most common emotions
        most_common_video = video_emotion_dist.most_common(1)[0] if video_emotion_dist else ('нет данных', 0)
        most_common_speech = speech_emotion_dist.most_common(1)[0] if speech_emotion_dist else ('нет данных', 0)
        
        # Calculate average confidence
        video_confidences = [item['confidence'] for item in video_emotions]
        speech_confidences = [item['confidence'] for item in speech_emotions]
        
        avg_video_conf = np.mean(video_confidences) if video_confidences else 0
        avg_speech_conf = np.mean(speech_confidences) if speech_confidences else 0
        
        # Generate HTML cards
        stats_html = f"""
        <div class="stat-card">
            <div class="stat-value">{metadata.get('duration', 0):.1f}с</div>
            <div>Общая продолжительность</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{len(video_emotions)}</div>
            <div>Анализов лица</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{len(speech_emotions)}</div>
            <div>Сегментов речи</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{most_common_video[0]}</div>
            <div>Частая эмоция (лицо)</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{most_common_speech[0]}</div>
            <div>Частая эмоция (речь)</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{avg_video_conf:.2f}</div>
            <div>Средняя уверенность (лицо)</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{avg_speech_conf:.2f}</div>
            <div>Средняя уверенность (речь)</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{len(data.get('critical_moments', []))}</div>
            <div>Критических моментов</div>
        </div>
        """
        
        return stats_html
    
    def _generate_critical_moments_html(self, data: Dict[str, Any]) -> str:
        """Generate HTML for critical moments"""
        critical_moments = data.get('critical_moments', [])
        
        if not critical_moments:
            return '<p>Критические моменты не обнаружены или не анализировались.</p>'
        
        moments_html = ''
        for moment in sorted(critical_moments, key=lambda x: x.get('timestamp', 0)):
            timestamp = moment.get('timestamp', 0)
            time_str = f"{int(timestamp//60)}:{int(timestamp%60):02d}"
            importance = moment.get('importance', 5)
            description = moment.get('description', 'Описание отсутствует')
            context = moment.get('context', '')
            emotions = moment.get('emotions', [])
            
            emotion_tags = ''.join([f'<span class="emotion-tag">{emotion}</span>' for emotion in emotions])
            
            moments_html += f"""
            <div class="critical-moment">
                <h4>⚠️ {time_str} - Важность: {importance}/10</h4>
                <p><strong>Описание:</strong> {description}</p>
                {f'<p><strong>Контекст:</strong> {context}</p>' if context else ''}
                {f'<p><strong>Связанные эмоции:</strong> {emotion_tags}</p>' if emotions else ''}
            </div>
            """
        
        return moments_html
    
    def _generate_transcript_html(self, data: Dict[str, Any]) -> str:
        """Generate HTML for transcript with emotion analysis"""
        speech_emotions = data.get('speech_emotions', [])
        transcript = data.get('transcript', [])
        
        if not speech_emotions and not transcript:
            return '<p>Транскрипция недоступна.</p>'
        
        transcript_html = ''
        
        # Use speech emotions if available, otherwise fallback to transcript
        segments = speech_emotions if speech_emotions else transcript
        
        for segment in segments:
            timestamp = segment.get('timestamp', 0)
            time_str = f"{int(timestamp//60)}:{int(timestamp%60):02d}"
            text = segment.get('text', 'Текст недоступен')
            speaker = segment.get('speaker', 'Неизвестный')
            emotion = segment.get('emotion', 'неопределено')
            confidence = segment.get('confidence', 0)
            
            # Get emotion color
            emotion_color = self.emotion_translator.get_emotion_color(emotion)
            
            transcript_html += f"""
            <div class="transcript-segment">
                <div class="timestamp">{time_str} - {speaker}</div>
                <p><strong>{text}</strong></p>
                <div style="margin-top: 10px;">
                    <span class="emotion-tag" style="background-color: {emotion_color}">
                        {emotion} ({confidence:.2f})
                    </span>
                </div>
            </div>
            """
        
        return transcript_html
    
    def _generate_insights_html(self, data: Dict[str, Any]) -> str:
        """Generate HTML for psychological insights"""
        insights = data.get('openai_insights', {})
        
        if not insights:
            return '<p>Психологические инсайты недоступны. Требуется интеграция с OpenAI.</p>'
        
        insights_html = ''
        
        # General analysis
        if 'general_analysis' in insights:
            insights_html += f"""
            <h3>🎯 Общий анализ</h3>
            <p>{insights['general_analysis']}</p>
            """
        
        # Emotional patterns
        if 'emotional_patterns' in insights:
            insights_html += f"""
            <h3>📈 Эмоциональные паттерны</h3>
            <p>{insights['emotional_patterns']}</p>
            """
        
        # Behavioral indicators
        if 'behavioral_indicators' in insights:
            insights_html += f"""
            <h3>🎭 Поведенческие индикаторы</h3>
            <p>{insights['behavioral_indicators']}</p>
            """
        
        # Recommendations
        if 'recommendations' in insights:
            insights_html += f"""
            <h3>💡 Рекомендации</h3>
            <p>{insights['recommendations']}</p>
            """
        
        # Risk assessment
        if 'risk_assessment' in insights:
            insights_html += f"""
            <h3>⚡ Оценка рисков</h3>
            <p>{insights['risk_assessment']}</p>
            """
        
        return insights_html if insights_html else '<p>Инсайты обрабатываются...</p>'
    
    def generate_json_report(self, data: Dict[str, Any], report_id: str) -> str:
        """Generate comprehensive JSON report"""
        output_path = str(self.reports_dir / f"{report_id}_report.json")
        
        json_data = {
            'metadata': {
                'report_id': report_id,
                'generated_at': datetime.now().isoformat(),
                'version': 'ДОПРОС MVP 2.0',
                'duration': data.get('metadata', {}).get('duration', 0),
                'video_frames': data.get('metadata', {}).get('video_frames', 0),
                'audio_segments': data.get('metadata', {}).get('audio_segments', 0)
            },
            'video_emotions': data.get('video_emotions', []),
            'speech_emotions': data.get('speech_emotions', []),
            'critical_moments': data.get('critical_moments', []),
            'transcript': data.get('transcript', []),
            'correlations': data.get('correlations', {}),
            'statistics': data.get('statistics', {}),
            'openai_insights': data.get('openai_insights', {})
        }
        
        # Add computed statistics
        video_emotions = data.get('video_emotions', [])
        speech_emotions = data.get('speech_emotions', [])
        
        if video_emotions:
            json_data['computed_statistics'] = {
                'video_emotion_distribution': dict(Counter([item['emotion'] for item in video_emotions])),
                'video_avg_confidence': float(np.mean([item['confidence'] for item in video_emotions])),
                'video_emotion_changes': self._count_emotion_changes(video_emotions)
            }
        
        if speech_emotions:
            json_data['computed_statistics'].update({
                'speech_emotion_distribution': dict(Counter([item['emotion'] for item in speech_emotions])),
                'speech_avg_confidence': float(np.mean([item['confidence'] for item in speech_emotions])),
                'speech_emotion_changes': self._count_emotion_changes(speech_emotions)
            })
        
        # Write JSON file
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, ensure_ascii=False, indent=2)
        
        return output_path
    
    def _create_critical_moments_report(self, data: Dict[str, Any], report_id: str) -> str:
        """Create critical moments CSV"""
        
        output_path = str(self.reports_dir / f"{report_id}_critical_moments.csv")
        
        critical_data = []
        for moment in data.get('critical_moments', []):
            critical_data.append({
                'timestamp': moment.get('timestamp', 0),
                'type': moment.get('type', ''),
                'description': moment.get('description', ''),
                'importance': moment.get('importance', 0),
                'confidence': moment.get('confidence', 0),
                'emotions_involved': ', '.join(moment.get('emotions', [])),
                'context': moment.get('context', '')
            })
        
        # Write to CSV
        if critical_data:
            df = pd.DataFrame(critical_data)
            df.to_csv(output_path, index=False, encoding='utf-8')
        
        return output_path
    
    def _generate_summary_statistics(self, data: Dict[str, Any]) -> str:
        """Generate HTML for summary statistics"""
        metadata = data.get('metadata', {})
        
        # Calculate emotion distribution
        video_emotions = data.get('video_emotions', [])
        speech_emotions = data.get('speech_emotions', [])
        
        video_emotion_dist = Counter([item['emotion'] for item in video_emotions])
        speech_emotion_dist = Counter([item['emotion'] for item in speech_emotions])
        
        # Most common emotions
        most_common_video = video_emotion_dist.most_common(1)[0] if video_emotion_dist else ('нет данных', 0)
        most_common_speech = speech_emotion_dist.most_common(1)[0] if speech_emotion_dist else ('нет данных', 0)
        
        # Calculate average confidence
        video_confidences = [item['confidence'] for item in video_emotions]
        speech_confidences = [item['confidence'] for item in speech_emotions]
        
        avg_video_conf = np.mean(video_confidences) if video_confidences else 0
        avg_speech_conf = np.mean(speech_confidences) if speech_confidences else 0
        
        # Generate HTML cards
        stats_html = f"""
        <div class="stat-card">
            <div class="stat-value">{metadata.get('duration', 0):.1f}с</div>
            <div>Общая продолжительность</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{len(video_emotions)}</div>
            <div>Анализов лица</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{len(speech_emotions)}</div>
            <div>Сегментов речи</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{most_common_video[0]}</div>
            <div>Частая эмоция (лицо)</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{most_common_speech[0]}</div>
            <div>Частая эмоция (речь)</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{avg_video_conf:.2f}</div>
            <div>Средняя уверенность (лицо)</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{avg_speech_conf:.2f}</div>
            <div>Средняя уверенность (речь)</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{len(data.get('critical_moments', []))}</div>
            <div>Критических моментов</div>
        </div>
        """
        
        return stats_html
    
    def _generate_critical_moments_html(self, data: Dict[str, Any]) -> str:
        """Generate HTML for critical moments"""
        critical_moments = data.get('critical_moments', [])
        
        if not critical_moments:
            return '<p>Критические моменты не обнаружены или не анализировались.</p>'
        
        moments_html = ''
        for moment in sorted(critical_moments, key=lambda x: x.get('timestamp', 0)):
            timestamp = moment.get('timestamp', 0)
            time_str = f"{int(timestamp//60)}:{int(timestamp%60):02d}"
            importance = moment.get('importance', 5)
            description = moment.get('description', 'Описание отсутствует')
            context = moment.get('context', '')
            emotions = moment.get('emotions', [])
            
            emotion_tags = ''.join([f'<span class="emotion-tag">{emotion}</span>' for emotion in emotions])
            
            moments_html += f"""
            <div class="critical-moment">
                <h4>⚠️ {time_str} - Важность: {importance}/10</h4>
                <p><strong>Описание:</strong> {description}</p>
                {f'<p><strong>Контекст:</strong> {context}</p>' if context else ''}
                {f'<p><strong>Связанные эмоции:</strong> {emotion_tags}</p>' if emotions else ''}
            </div>
            """
        
        return moments_html
    
    def _generate_transcript_html(self, data: Dict[str, Any]) -> str:
        """Generate HTML for transcript with emotion analysis"""
        speech_emotions = data.get('speech_emotions', [])
        transcript = data.get('transcript', [])
        
        if not speech_emotions and not transcript:
            return '<p>Транскрипция недоступна.</p>'
        
        transcript_html = ''
        
        # Use speech emotions if available, otherwise fallback to transcript
        segments = speech_emotions if speech_emotions else transcript
        
        for segment in segments:
            timestamp = segment.get('timestamp', 0)
            time_str = f"{int(timestamp//60)}:{int(timestamp%60):02d}"
            text = segment.get('text', 'Текст недоступен')
            speaker = segment.get('speaker', 'Неизвестный')
            emotion = segment.get('emotion', 'неопределено')
            confidence = segment.get('confidence', 0)
            
            # Get emotion color
            emotion_color = self.emotion_translator.get_emotion_color(emotion)
            
            transcript_html += f"""
            <div class="transcript-segment">
                <div class="timestamp">{time_str} - {speaker}</div>
                <p><strong>{text}</strong></p>
                <div style="margin-top: 10px;">
                    <span class="emotion-tag" style="background-color: {emotion_color}">
                        {emotion} ({confidence:.2f})
                    </span>
                </div>
            </div>
            """
        
        return transcript_html
    
    def _generate_insights_html(self, data: Dict[str, Any]) -> str:
        """Generate HTML for psychological insights"""
        insights = data.get('openai_insights', {})
        
        if not insights:
            return '<p>Психологические инсайты недоступны. Требуется интеграция с OpenAI.</p>'
        
        insights_html = ''
        
        # General analysis
        if 'general_analysis' in insights:
            insights_html += f"""
            <h3>🎯 Общий анализ</h3>
            <p>{insights['general_analysis']}</p>
            """
        
        # Emotional patterns
        if 'emotional_patterns' in insights:
            insights_html += f"""
            <h3>📈 Эмоциональные паттерны</h3>
            <p>{insights['emotional_patterns']}</p>
            """
        
        # Behavioral indicators
        if 'behavioral_indicators' in insights:
            insights_html += f"""
            <h3>🎭 Поведенческие индикаторы</h3>
            <p>{insights['behavioral_indicators']}</p>
            """
        
        # Recommendations
        if 'recommendations' in insights:
            insights_html += f"""
            <h3>💡 Рекомендации</h3>
            <p>{insights['recommendations']}</p>
            """
        
        # Risk assessment
        if 'risk_assessment' in insights:
            insights_html += f"""
            <h3>⚡ Оценка рисков</h3>
            <p>{insights['risk_assessment']}</p>
            """
        
        return insights_html if insights_html else '<p>Инсайты обрабатываются...</p>'
    
    def generate_json_report(self, data: Dict[str, Any], report_id: str) -> str:
        """Generate comprehensive JSON report"""
        output_path = str(self.reports_dir / f"{report_id}_report.json")
        
        json_data = {
            'metadata': {
                'report_id': report_id,
                'generated_at': datetime.now().isoformat(),
                'version': 'ДОПРОС MVP 2.0',
                'duration': data.get('metadata', {}).get('duration', 0),
                'video_frames': data.get('metadata', {}).get('video_frames', 0),
                'audio_segments': data.get('metadata', {}).get('audio_segments', 0)
            },
            'video_emotions': data.get('video_emotions', []),
            'speech_emotions': data.get('speech_emotions', []),
            'critical_moments': data.get('critical_moments', []),
            'transcript': data.get('transcript', []),
            'correlations': data.get('correlations', {}),
            'statistics': data.get('statistics', {}),
            'openai_insights': data.get('openai_insights', {})
        }
        
        # Add computed statistics
        video_emotions = data.get('video_emotions', [])
        speech_emotions = data.get('speech_emotions', [])
        
        if video_emotions:
            json_data['computed_statistics'] = {
                'video_emotion_distribution': dict(Counter([item['emotion'] for item in video_emotions])),
                'video_avg_confidence': float(np.mean([item['confidence'] for item in video_emotions])),
                'video_emotion_changes': self._count_emotion_changes(video_emotions)
            }
        
        if speech_emotions:
            json_data['computed_statistics'].update({
                'speech_emotion_distribution': dict(Counter([item['emotion'] for item in speech_emotions])),
                'speech_avg_confidence': float(np.mean([item['confidence'] for item in speech_emotions])),
                'speech_emotion_changes': self._count_emotion_changes(speech_emotions)
            })
        
        # Write JSON file
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, ensure_ascii=False, indent=2)
        
        return output_path
    
    def _create_correlation_matrix(self, data: Dict[str, Any], report_id: str) -> str:
        """Create correlation matrix CSV"""
        
        output_path = str(self.reports_dir / f"{report_id}_correlation_matrix.csv")
        
        correlations = data.get('correlations', {})
        
        # Convert correlations to matrix format
        correlation_data = []
        for key, value in correlations.items():
            correlation_data.append({
                'metric_pair': key,
                'correlation_coefficient': value.get('correlation', 0),
                'p_value': value.get('p_value', 1.0),
                'significance': 'significant' if value.get('p_value', 1.0) < 0.05 else 'not_significant',
                'sample_size': value.get('sample_size', 0)
            })
        
        # Write to CSV
        if correlation_data:
            df = pd.DataFrame(correlation_data)
            df.to_csv(output_path, index=False, encoding='utf-8')
        
        return output_path
    
    def _generate_summary_statistics(self, data: Dict[str, Any]) -> str:
        """Generate HTML for summary statistics"""
        metadata = data.get('metadata', {})
        
        # Calculate emotion distribution
        video_emotions = data.get('video_emotions', [])
        speech_emotions = data.get('speech_emotions', [])
        
        video_emotion_dist = Counter([item['emotion'] for item in video_emotions])
        speech_emotion_dist = Counter([item['emotion'] for item in speech_emotions])
        
        # Most common emotions
        most_common_video = video_emotion_dist.most_common(1)[0] if video_emotion_dist else ('нет данных', 0)
        most_common_speech = speech_emotion_dist.most_common(1)[0] if speech_emotion_dist else ('нет данных', 0)
        
        # Calculate average confidence
        video_confidences = [item['confidence'] for item in video_emotions]
        speech_confidences = [item['confidence'] for item in speech_emotions]
        
        avg_video_conf = np.mean(video_confidences) if video_confidences else 0
        avg_speech_conf = np.mean(speech_confidences) if speech_confidences else 0
        
        # Generate HTML cards
        stats_html = f"""
        <div class="stat-card">
            <div class="stat-value">{metadata.get('duration', 0):.1f}с</div>
            <div>Общая продолжительность</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{len(video_emotions)}</div>
            <div>Анализов лица</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{len(speech_emotions)}</div>
            <div>Сегментов речи</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{most_common_video[0]}</div>
            <div>Частая эмоция (лицо)</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{most_common_speech[0]}</div>
            <div>Частая эмоция (речь)</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{avg_video_conf:.2f}</div>
            <div>Средняя уверенность (лицо)</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{avg_speech_conf:.2f}</div>
            <div>Средняя уверенность (речь)</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{len(data.get('critical_moments', []))}</div>
            <div>Критических моментов</div>
        </div>
        """
        
        return stats_html
    
    def _generate_critical_moments_html(self, data: Dict[str, Any]) -> str:
        """Generate HTML for critical moments"""
        critical_moments = data.get('critical_moments', [])
        
        if not critical_moments:
            return '<p>Критические моменты не обнаружены или не анализировались.</p>'
        
        moments_html = ''
        for moment in sorted(critical_moments, key=lambda x: x.get('timestamp', 0)):
            timestamp = moment.get('timestamp', 0)
            time_str = f"{int(timestamp//60)}:{int(timestamp%60):02d}"
            importance = moment.get('importance', 5)
            description = moment.get('description', 'Описание отсутствует')
            context = moment.get('context', '')
            emotions = moment.get('emotions', [])
            
            emotion_tags = ''.join([f'<span class="emotion-tag">{emotion}</span>' for emotion in emotions])
            
            moments_html += f"""
            <div class="critical-moment">
                <h4>⚠️ {time_str} - Важность: {importance}/10</h4>
                <p><strong>Описание:</strong> {description}</p>
                {f'<p><strong>Контекст:</strong> {context}</p>' if context else ''}
                {f'<p><strong>Связанные эмоции:</strong> {emotion_tags}</p>' if emotions else ''}
            </div>
            """
        
        return moments_html
    
    def _generate_transcript_html(self, data: Dict[str, Any]) -> str:
        """Generate HTML for transcript with emotion analysis"""
        speech_emotions = data.get('speech_emotions', [])
        transcript = data.get('transcript', [])
        
        if not speech_emotions and not transcript:
            return '<p>Транскрипция недоступна.</p>'
        
        transcript_html = ''
        
        # Use speech emotions if available, otherwise fallback to transcript
        segments = speech_emotions if speech_emotions else transcript
        
        for segment in segments:
            timestamp = segment.get('timestamp', 0)
            time_str = f"{int(timestamp//60)}:{int(timestamp%60):02d}"
            text = segment.get('text', 'Текст недоступен')
            speaker = segment.get('speaker', 'Неизвестный')
            emotion = segment.get('emotion', 'неопределено')
            confidence = segment.get('confidence', 0)
            
            # Get emotion color
            emotion_color = self.emotion_translator.get_emotion_color(emotion)
            
            transcript_html += f"""
            <div class="transcript-segment">
                <div class="timestamp">{time_str} - {speaker}</div>
                <p><strong>{text}</strong></p>
                <div style="margin-top: 10px;">
                    <span class="emotion-tag" style="background-color: {emotion_color}">
                        {emotion} ({confidence:.2f})
                    </span>
                </div>
            </div>
            """
        
        return transcript_html
    
    def _generate_insights_html(self, data: Dict[str, Any]) -> str:
        """Generate HTML for psychological insights"""
        insights = data.get('openai_insights', {})
        
        if not insights:
            return '<p>Психологические инсайты недоступны. Требуется интеграция с OpenAI.</p>'
        
        insights_html = ''
        
        # General analysis
        if 'general_analysis' in insights:
            insights_html += f"""
            <h3>🎯 Общий анализ</h3>
            <p>{insights['general_analysis']}</p>
            """
        
        # Emotional patterns
        if 'emotional_patterns' in insights:
            insights_html += f"""
            <h3>📈 Эмоциональные паттерны</h3>
            <p>{insights['emotional_patterns']}</p>
            """
        
        # Behavioral indicators
        if 'behavioral_indicators' in insights:
            insights_html += f"""
            <h3>🎭 Поведенческие индикаторы</h3>
            <p>{insights['behavioral_indicators']}</p>
            """
        
        # Recommendations
        if 'recommendations' in insights:
            insights_html += f"""
            <h3>💡 Рекомендации</h3>
            <p>{insights['recommendations']}</p>
            """
        
        # Risk assessment
        if 'risk_assessment' in insights:
            insights_html += f"""
            <h3>⚡ Оценка рисков</h3>
            <p>{insights['risk_assessment']}</p>
            """
        
        return insights_html if insights_html else '<p>Инсайты обрабатываются...</p>'
    
    def generate_json_report(self, data: Dict[str, Any], report_id: str) -> str:
        """Generate comprehensive JSON report"""
        output_path = str(self.reports_dir / f"{report_id}_report.json")
        
        json_data = {
            'metadata': {
                'report_id': report_id,
                'generated_at': datetime.now().isoformat(),
                'version': 'ДОПРОС MVP 2.0',
                'duration': data.get('metadata', {}).get('duration', 0),
                'video_frames': data.get('metadata', {}).get('video_frames', 0),
                'audio_segments': data.get('metadata', {}).get('audio_segments', 0)
            },
            'video_emotions': data.get('video_emotions', []),
            'speech_emotions': data.get('speech_emotions', []),
            'critical_moments': data.get('critical_moments', []),
            'transcript': data.get('transcript', []),
            'correlations': data.get('correlations', {}),
            'statistics': data.get('statistics', {}),
            'openai_insights': data.get('openai_insights', {})
        }
        
        # Add computed statistics
        video_emotions = data.get('video_emotions', [])
        speech_emotions = data.get('speech_emotions', [])
        
        if video_emotions:
            json_data['computed_statistics'] = {
                'video_emotion_distribution': dict(Counter([item['emotion'] for item in video_emotions])),
                'video_avg_confidence': float(np.mean([item['confidence'] for item in video_emotions])),
                'video_emotion_changes': self._count_emotion_changes(video_emotions)
            }
        
        if speech_emotions:
            json_data['computed_statistics'].update({
                'speech_emotion_distribution': dict(Counter([item['emotion'] for item in speech_emotions])),
                'speech_avg_confidence': float(np.mean([item['confidence'] for item in speech_emotions])),
                'speech_emotion_changes': self._count_emotion_changes(speech_emotions)
            })
        
        # Write JSON file
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, ensure_ascii=False, indent=2)
        
        return output_path
    
    def _create_statistics_report(self, data: Dict[str, Any], report_id: str) -> str:
        """Create comprehensive statistics CSV"""
        
        output_path = str(self.reports_dir / f"{report_id}_summary_statistics.csv")
        
        stats = data.get('statistics', {})
        
        stats_data = []
        
        # Overall statistics
        stats_data.append({
            'category': 'overall',
            'metric': 'total_duration',
            'value': data.get('metadata', {}).get('duration', 0),
            'unit': 'seconds'
        })
        
        stats_data.append({
            'category': 'overall',
            'metric': 'video_frames',
            'value': data.get('metadata', {}).get('video_frames', 0),
            'unit': 'frames'
        })
        
        stats_data.append({
            'category': 'overall',
            'metric': 'audio_segments',
            'value': data.get('metadata', {}).get('audio_segments', 0),
            'unit': 'segments'
        })
        
        # Emotion distribution statistics
        for source in ['video', 'speech']:
            emotion_stats = stats.get(f'{source}_emotion_distribution', {})
            for emotion, count in emotion_stats.items():
                stats_data.append({
                    'category': f'{source}_emotions',
                    'metric': emotion,
                    'value': count,
                    'unit': 'occurrences'
                })
        
        # Write to CSV
        if stats_data:
            df = pd.DataFrame(stats_data)
            df.to_csv(output_path, index=False, encoding='utf-8')
        
        return output_path
    
    def _generate_summary_statistics(self, data: Dict[str, Any]) -> str:
        """Generate HTML for summary statistics"""
        metadata = data.get('metadata', {})
        
        # Calculate emotion distribution
        video_emotions = data.get('video_emotions', [])
        speech_emotions = data.get('speech_emotions', [])
        
        video_emotion_dist = Counter([item['emotion'] for item in video_emotions])
        speech_emotion_dist = Counter([item['emotion'] for item in speech_emotions])
        
        # Most common emotions
        most_common_video = video_emotion_dist.most_common(1)[0] if video_emotion_dist else ('нет данных', 0)
        most_common_speech = speech_emotion_dist.most_common(1)[0] if speech_emotion_dist else ('нет данных', 0)
        
        # Calculate average confidence
        video_confidences = [item['confidence'] for item in video_emotions]
        speech_confidences = [item['confidence'] for item in speech_emotions]
        
        avg_video_conf = np.mean(video_confidences) if video_confidences else 0
        avg_speech_conf = np.mean(speech_confidences) if speech_confidences else 0
        
        # Generate HTML cards
        stats_html = f"""
        <div class="stat-card">
            <div class="stat-value">{metadata.get('duration', 0):.1f}с</div>
            <div>Общая продолжительность</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{len(video_emotions)}</div>
            <div>Анализов лица</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{len(speech_emotions)}</div>
            <div>Сегментов речи</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{most_common_video[0]}</div>
            <div>Частая эмоция (лицо)</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{most_common_speech[0]}</div>
            <div>Частая эмоция (речь)</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{avg_video_conf:.2f}</div>
            <div>Средняя уверенность (лицо)</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{avg_speech_conf:.2f}</div>
            <div>Средняя уверенность (речь)</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{len(data.get('critical_moments', []))}</div>
            <div>Критических моментов</div>
        </div>
        """
        
        return stats_html
    
    def _generate_critical_moments_html(self, data: Dict[str, Any]) -> str:
        """Generate HTML for critical moments"""
        critical_moments = data.get('critical_moments', [])
        
        if not critical_moments:
            return '<p>Критические моменты не обнаружены или не анализировались.</p>'
        
        moments_html = ''
        for moment in sorted(critical_moments, key=lambda x: x.get('timestamp', 0)):
            timestamp = moment.get('timestamp', 0)
            time_str = f"{int(timestamp//60)}:{int(timestamp%60):02d}"
            importance = moment.get('importance', 5)
            description = moment.get('description', 'Описание отсутствует')
            context = moment.get('context', '')
            emotions = moment.get('emotions', [])
            
            emotion_tags = ''.join([f'<span class="emotion-tag">{emotion}</span>' for emotion in emotions])
            
            moments_html += f"""
            <div class="critical-moment">
                <h4>⚠️ {time_str} - Важность: {importance}/10</h4>
                <p><strong>Описание:</strong> {description}</p>
                {f'<p><strong>Контекст:</strong> {context}</p>' if context else ''}
                {f'<p><strong>Связанные эмоции:</strong> {emotion_tags}</p>' if emotions else ''}
            </div>
            """
        
        return moments_html
    
    def _generate_transcript_html(self, data: Dict[str, Any]) -> str:
        """Generate HTML for transcript with emotion analysis"""
        speech_emotions = data.get('speech_emotions', [])
        transcript = data.get('transcript', [])
        
        if not speech_emotions and not transcript:
            return '<p>Транскрипция недоступна.</p>'
        
        transcript_html = ''
        
        # Use speech emotions if available, otherwise fallback to transcript
        segments = speech_emotions if speech_emotions else transcript
        
        for segment in segments:
            timestamp = segment.get('timestamp', 0)
            time_str = f"{int(timestamp//60)}:{int(timestamp%60):02d}"
            text = segment.get('text', 'Текст недоступен')
            speaker = segment.get('speaker', 'Неизвестный')
            emotion = segment.get('emotion', 'неопределено')
            confidence = segment.get('confidence', 0)
            
            # Get emotion color
            emotion_color = self.emotion_translator.get_emotion_color(emotion)
            
            transcript_html += f"""
            <div class="transcript-segment">
                <div class="timestamp">{time_str} - {speaker}</div>
                <p><strong>{text}</strong></p>
                <div style="margin-top: 10px;">
                    <span class="emotion-tag" style="background-color: {emotion_color}">
                        {emotion} ({confidence:.2f})
                    </span>
                </div>
            </div>
            """
        
        return transcript_html
    
    def _generate_insights_html(self, data: Dict[str, Any]) -> str:
        """Generate HTML for psychological insights"""
        insights = data.get('openai_insights', {})
        
        if not insights:
            return '<p>Психологические инсайты недоступны. Требуется интеграция с OpenAI.</p>'
        
        insights_html = ''
        
        # General analysis
        if 'general_analysis' in insights:
            insights_html += f"""
            <h3>🎯 Общий анализ</h3>
            <p>{insights['general_analysis']}</p>
            """
        
        # Emotional patterns
        if 'emotional_patterns' in insights:
            insights_html += f"""
            <h3>📈 Эмоциональные паттерны</h3>
            <p>{insights['emotional_patterns']}</p>
            """
        
        # Behavioral indicators
        if 'behavioral_indicators' in insights:
            insights_html += f"""
            <h3>🎭 Поведенческие индикаторы</h3>
            <p>{insights['behavioral_indicators']}</p>
            """
        
        # Recommendations
        if 'recommendations' in insights:
            insights_html += f"""
            <h3>💡 Рекомендации</h3>
            <p>{insights['recommendations']}</p>
            """
        
        # Risk assessment
        if 'risk_assessment' in insights:
            insights_html += f"""
            <h3>⚡ Оценка рисков</h3>
            <p>{insights['risk_assessment']}</p>
            """
        
        return insights_html if insights_html else '<p>Инсайты обрабатываются...</p>'
    
    def generate_json_report(self, data: Dict[str, Any], report_id: str) -> str:
        """Generate comprehensive JSON report"""
        output_path = str(self.reports_dir / f"{report_id}_report.json")
        
        json_data = {
            'metadata': {
                'report_id': report_id,
                'generated_at': datetime.now().isoformat(),
                'version': 'ДОПРОС MVP 2.0',
                'duration': data.get('metadata', {}).get('duration', 0),
                'video_frames': data.get('metadata', {}).get('video_frames', 0),
                'audio_segments': data.get('metadata', {}).get('audio_segments', 0)
            },
            'video_emotions': data.get('video_emotions', []),
            'speech_emotions': data.get('speech_emotions', []),
            'critical_moments': data.get('critical_moments', []),
            'transcript': data.get('transcript', []),
            'correlations': data.get('correlations', {}),
            'statistics': data.get('statistics', {}),
            'openai_insights': data.get('openai_insights', {})
        }
        
        # Add computed statistics
        video_emotions = data.get('video_emotions', [])
        speech_emotions = data.get('speech_emotions', [])
        
        if video_emotions:
            json_data['computed_statistics'] = {
                'video_emotion_distribution': dict(Counter([item['emotion'] for item in video_emotions])),
                'video_avg_confidence': float(np.mean([item['confidence'] for item in video_emotions])),
                'video_emotion_changes': self._count_emotion_changes(video_emotions)
            }
        
        if speech_emotions:
            json_data['computed_statistics'].update({
                'speech_emotion_distribution': dict(Counter([item['emotion'] for item in speech_emotions])),
                'speech_avg_confidence': float(np.mean([item['confidence'] for item in speech_emotions])),
                'speech_emotion_changes': self._count_emotion_changes(speech_emotions)
            })
        
        # Write JSON file
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, ensure_ascii=False, indent=2)
        
        return output_path
    
    def _generate_all_visualizations(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate all interactive visualizations"""
        
        visualizations = {}
        
        try:
            # 1. Interactive timeline
            timeline_fig = self.create_interactive_timeline(data)
            visualizations['timeline'] = self._fig_to_dict(timeline_fig)
            
            # 2. Emotion heatmap
            heatmap_fig = self.create_emotion_heatmap(data)
            visualizations['heatmap'] = self._fig_to_dict(heatmap_fig)
            
            # 3. Statistical charts
            stats_figs = self.create_statistical_charts(data)
            visualizations['statistics'] = {
                name: self._fig_to_dict(fig) 
                for name, fig in stats_figs.items()
            }
            
            # 4. Correlation plot
            correlation_fig = self.create_correlation_plot(data)
            visualizations['correlations'] = self._fig_to_dict(correlation_fig)
            
            # 5. Critical moments visualization
            critical_fig = self.create_critical_moments_plot(data)
            visualizations['critical_moments'] = self._fig_to_dict(critical_fig)
            
        except Exception as e:
            self.logger.error(f"Visualization generation failed: {e}")
        
        return visualizations
    
    def create_interactive_timeline(self, data: Dict[str, Any]) -> go.Figure:
        """Create interactive timeline with video and speech emotions"""
        
        fig = make_subplots(
            rows=2, cols=1,
            subplot_titles=('Эмоции лица (видео)', 'Эмоции речи (аудио)'),
            vertical_spacing=0.15
        )
        
        # Video emotions timeline
        video_emotions = data.get('video_emotions', [])
        if video_emotions:
            video_timestamps = [item['timestamp'] for item in video_emotions]
            video_emotion_names = [item['emotion'] for item in video_emotions]
            video_confidences = [item['confidence'] for item in video_emotions]
            
            # Convert emotions to numeric values for plotting
            unique_emotions = list(set(video_emotion_names))
            emotion_to_num = {emotion: i for i, emotion in enumerate(unique_emotions)}
            video_emotion_nums = [emotion_to_num[emotion] for emotion in video_emotion_names]
            
            fig.add_trace(
                go.Scatter(
                    x=video_timestamps,
                    y=video_emotion_nums,
                    mode='lines+markers',
                    name='Эмоции лица',
                    line=dict(color='blue', width=2),
                    marker=dict(size=8),
                    customdata=list(zip(video_emotion_names, video_confidences)),
                    hovertemplate='<b>%{customdata[0]}</b><br>' +
                                'Время: %{x:.1f}с<br>' +
                                'Уверенность: %{customdata[1]:.2f}<br>' +
                                '<extra></extra>'
                ),
                row=1, col=1
            )
            
            # Update y-axis labels for video emotions
            fig.update_yaxes(
                tickmode='array',
                tickvals=list(range(len(unique_emotions))),
                ticktext=unique_emotions,
                row=1, col=1
            )
        
        # Speech emotions timeline
        speech_emotions = data.get('speech_emotions', [])
        if speech_emotions:
            speech_timestamps = [item['timestamp'] for item in speech_emotions]
            speech_emotion_names = [item['emotion'] for item in speech_emotions]
            speech_confidences = [item['confidence'] for item in speech_emotions]
            speech_texts = [item.get('text', '')[:50] + '...' if len(item.get('text', '')) > 50 
                          else item.get('text', '') for item in speech_emotions]
            
            # Convert emotions to numeric values
            unique_speech_emotions = list(set(speech_emotion_names))
            speech_emotion_to_num = {emotion: i for i, emotion in enumerate(unique_speech_emotions)}
            speech_emotion_nums = [speech_emotion_to_num[emotion] for emotion in speech_emotion_names]
            
            fig.add_trace(
                go.Scatter(
                    x=speech_timestamps,
                    y=speech_emotion_nums,
                    mode='lines+markers',
                    name='Эмоции речи',
                    line=dict(color='red', width=2),
                    marker=dict(size=8),
                    customdata=list(zip(speech_emotion_names, speech_confidences, speech_texts)),
                    hovertemplate='<b>%{customdata[0]}</b><br>' +
                                'Время: %{x:.1f}с<br>' +
                                'Уверенность: %{customdata[1]:.2f}<br>' +
                                'Текст: %{customdata[2]}<br>' +
                                '<extra></extra>'
                ),
                row=2, col=1
            )
            
            # Update y-axis labels for speech emotions
            fig.update_yaxes(
                tickmode='array',
                tickvals=list(range(len(unique_speech_emotions))),
                ticktext=unique_speech_emotions,
                row=2, col=1
            )
        
        # Add critical moments as vertical lines
        for moment in data.get('critical_moments', []):
            timestamp = moment.get('timestamp', 0)
            description = moment.get('description', '')
            
            fig.add_vline(
                x=timestamp,
                line_dash="dash",
                line_color="orange",
                annotation_text=description[:30] + '...' if len(description) > 30 else description,
                annotation_position="top"
            )
        
        # Update layout
        fig.update_layout(
            title="Временная шкала эмоций",
            height=800,
            showlegend=True,
            hovermode='x unified'
        )
        
        fig.update_xaxes(title_text="Время (секунды)")
        
        return fig
    
    def create_emotion_heatmap(self, data: Dict[str, Any]) -> go.Figure:
        """Create emotion intensity heatmap over time"""
        
        # Combine all emotion data
        all_emotions = []
        all_emotions.extend(data.get('video_emotions', []))
        all_emotions.extend(data.get('speech_emotions', []))
        
        if not all_emotions:
            # Return empty figure if no data
            fig = go.Figure()
            fig.update_layout(title="Тепловая карта эмоций (нет данных)")
            return fig
        
        # Get unique emotions and time bins
        unique_emotions = list(set([item['emotion'] for item in all_emotions]))
        
        # Create time bins (30-second intervals)
        max_time = max([item['timestamp'] for item in all_emotions])
        time_bins = np.arange(0, max_time + 30, 30)
        time_labels = [f"{int(t/60)}:{int(t%60):02d}" for t in time_bins[:-1]]
        
        # Create emotion intensity matrix
        emotion_matrix = np.zeros((len(unique_emotions), len(time_bins) - 1))
        
        for item in all_emotions:
            emotion_idx = unique_emotions.index(item['emotion'])
            time_bin_idx = np.digitize(item['timestamp'], time_bins) - 1
            if 0 <= time_bin_idx < len(time_bins) - 1:
                emotion_matrix[emotion_idx, time_bin_idx] += item['confidence']
        
        # Normalize by number of observations in each bin
        for j in range(emotion_matrix.shape[1]):
            bin_count = np.sum(emotion_matrix[:, j] > 0)
            if bin_count > 0:
                emotion_matrix[:, j] /= bin_count
        
        # Create heatmap
        fig = px.imshow(
            emotion_matrix,
            labels=dict(x="Время", y="Эмоция", color="Интенсивность"),
            x=time_labels,
            y=unique_emotions,
            color_continuous_scale="RdBu_r",
            title="Тепловая карта эмоциональной интенсивности"
        )
        
        fig.update_layout(
            height=600,
            xaxis_title="Время (мм:сс)",
            yaxis_title="Эмоции"
        )
        
        return fig
    
    def create_statistical_charts(self, data: Dict[str, Any]) -> Dict[str, go.Figure]:
        """Create various statistical charts"""
        
        charts = {}
        
        # 1. Emotion distribution pie charts
        for source in ['video', 'speech']:
            emotions_data = data.get(f'{source}_emotions', [])
            if emotions_data:
                emotion_counts = Counter([item['emotion'] for item in emotions_data])
                
                fig = go.Figure(data=[
                    go.Pie(
                        labels=list(emotion_counts.keys()),
                        values=list(emotion_counts.values()),
                        textinfo='label+percent',
                        insidetextorientation='radial'
                    )
                ])
                
                fig.update_layout(
                    title=f"Распределение эмоций ({'видео' if source == 'video' else 'речь'})"
                )
                
                charts[f'{source}_distribution'] = fig
        
        # 2. Confidence distribution histogram
        all_confidences = []
        sources = []
        
        for item in data.get('video_emotions', []):
            all_confidences.append(item['confidence'])
            sources.append('Видео')
        
        for item in data.get('speech_emotions', []):
            all_confidences.append(item['confidence'])
            sources.append('Речь')
        
        if all_confidences:
            fig = go.Figure()
            
            # Add histograms for each source
            for source in ['Видео', 'Речь']:
                source_confidences = [conf for conf, src in zip(all_confidences, sources) if src == source]
                if source_confidences:
                    fig.add_trace(go.Histogram(
                        x=source_confidences,
                        name=source,
                        opacity=0.7,
                        nbinsx=20
                    ))
            
            fig.update_layout(
                title="Распределение уверенности анализа",
                xaxis_title="Уверенность",
                yaxis_title="Частота",
                barmode='overlay'
            )
            
            charts['confidence_distribution'] = fig
        
        # 3. Emotion changes over time
        if data.get('video_emotions') and data.get('speech_emotions'):
            
            # Calculate emotion change rate
            def calculate_change_rate(emotions_list):
                changes = []
                for i in range(1, len(emotions_list)):
                    prev_emotion = emotions_list[i-1]['emotion']
                    curr_emotion = emotions_list[i]['emotion']
                    if prev_emotion != curr_emotion:
                        changes.append(emotions_list[i]['timestamp'])
                return changes
            
            video_changes = calculate_change_rate(data['video_emotions'])
            speech_changes = calculate_change_rate(data['speech_emotions'])
            
            # Create time bins and count changes
            max_time = max([
                data['video_emotions'][-1]['timestamp'] if data['video_emotions'] else 0,
                data['speech_emotions'][-1]['timestamp'] if data['speech_emotions'] else 0
            ])
            
            time_bins = np.arange(0, max_time + 60, 60)  # 1-minute bins
            time_labels = [f"{int(t/60)}" for t in time_bins[:-1]]
            
            video_change_counts = np.histogram(video_changes, bins=time_bins)[0]
            speech_change_counts = np.histogram(speech_changes, bins=time_bins)[0]
            
            fig = go.Figure()
            
            fig.add_trace(go.Bar(
                x=time_labels,
                y=video_change_counts,
                name='Изменения в видео',
                marker_color='blue'
            ))
            
            fig.add_trace(go.Bar(
                x=time_labels,
                y=speech_change_counts,
                name='Изменения в речи',
                marker_color='red'
            ))
            
            fig.update_layout(
                title="Частота изменений эмоций по времени",
                xaxis_title="Время (минуты)",
                yaxis_title="Количество изменений",
                barmode='group'
            )
            
            charts['emotion_changes'] = fig
        
        return charts
    
    def create_correlation_plot(self, data: Dict[str, Any]) -> go.Figure:
        """Create correlation visualization between video and speech emotions"""
        
        correlations = data.get('correlations', {})
        
        if not correlations:
            fig = go.Figure()
            fig.update_layout(title="Корреляции (нет данных)")
            return fig
        
        # Extract correlation data
        correlation_names = []
        correlation_values = []
        p_values = []
        
        for name, corr_data in correlations.items():
            correlation_names.append(name.replace('_', ' ').title())
            correlation_values.append(corr_data.get('correlation', 0))
            p_values.append(corr_data.get('p_value', 1.0))
        
        # Create color based on significance
        colors = ['red' if p < 0.05 else 'gray' for p in p_values]
        
        fig = go.Figure(data=[
            go.Bar(
                x=correlation_names,
                y=correlation_values,
                marker_color=colors,
                text=[f"{val:.3f}" for val in correlation_values],
                textposition='auto'
            )
        ])
        
        fig.update_layout(
            title="Корреляции между показателями",
            xaxis_title="Сравниваемые показатели",
            yaxis_title="Коэффициент корреляции",
            yaxis=dict(range=[-1, 1])
        )
        
        # Add horizontal line at y=0
        fig.add_hline(y=0, line_dash="dash", line_color="black")
        
        return fig
    
    def create_critical_moments_plot(self, data: Dict[str, Any]) -> go.Figure:
        """Create critical moments visualization"""
        
        critical_moments = data.get('critical_moments', [])
        
        if not critical_moments:
            fig = go.Figure()
            fig.update_layout(title="Критические моменты (не найдены)")
            return fig
        
        # Extract data
        timestamps = [moment.get('timestamp', 0) for moment in critical_moments]
        importance = [moment.get('importance', 5) for moment in critical_moments]
        types = [moment.get('type', 'unknown') for moment in critical_moments]
        descriptions = [moment.get('description', '') for moment in critical_moments]
        
        # Create scatter plot
        fig = go.Figure(data=go.Scatter(
            x=timestamps,
            y=importance,
            mode='markers',
            marker=dict(
                size=[imp * 5 for imp in importance],  # Size based on importance
                color=importance,
                colorscale='Reds',
                showscale=True,
                colorbar=dict(title="Важность")
            ),
            text=descriptions,
            hovertemplate='<b>%{text}</b><br>' +
                         'Время: %{x:.1f}с<br>' +
                         'Важность: %{y}/10<br>' +
                         '<extra></extra>'
        ))
        
        fig.update_layout(
            title="Критические моменты допроса",
            xaxis_title="Время (секунды)",
            yaxis_title="Важность (1-10)",
            height=500
        )
        
        return fig
    
    def _fig_to_dict(self, fig: go.Figure) -> Dict[str, Any]:
        """Convert plotly figure to dictionary for storage"""
        if not fig:
            return {}
        
        return {
            'data': fig.data,
            'layout': fig.layout,
            'html': fig.to_html(include_plotlyjs='cdn', div_id=None),
            'json': fig.to_json()
        }
    
    def generate_html_report(self, data: Dict[str, Any], visualizations: Dict[str, Any], 
                           report_id: str) -> str:
        """Generate comprehensive HTML report"""
        
        output_path = str(self.reports_dir / f"{report_id}_report.html")
        
        # Custom CSS
        custom_css = """
        body { 
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; 
            margin: 40px; 
            background-color: #f8f9fa;
        }
        .header {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 30px;
            border-radius: 10px;
            margin-bottom: 30px;
            text-align: center;
        }
        .section {
            background: white;
            padding: 25px;
            margin: 20px 0;
            border-radius: 8px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }
        .section h2 {
            color: #333;
            border-bottom: 3px solid #667eea;
            padding-bottom: 10px;
        }
        .stat-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 20px;
            margin: 20px 0;
        }
        .stat-card {
            background: #f8f9fa;
            padding: 20px;
            border-radius: 8px;
            text-align: center;
            border-left: 4px solid #667eea;
        }
        .stat-value {
            font-size: 2em;
            font-weight: bold;
            color: #667eea;
        }
        .critical-moment {
            background: #fff3cd;
            border: 1px solid #ffeeba;
            border-radius: 5px;
            padding: 15px;
            margin: 10px 0;
        }
        .transcript-segment {
            margin: 10px 0;
            padding: 15px;
            background: #f8f9fa;
            border-left: 4px solid #28a745;
            border-radius: 4px;
        }
        .timestamp {
            color: #6c757d;
            font-size: 0.9em;
        }
        .emotion-tag {
            display: inline-block;
            padding: 4px 8px;
            background: #667eea;
            color: white;
            border-radius: 12px;
            font-size: 0.8em;
            margin: 2px;
        }
        """
        
        # Generate summary statistics
        summary_stats = self._generate_summary_statistics(data)
        
        # Generate critical moments HTML
        critical_moments_html = self._generate_critical_moments_html(data)
        
        # Generate transcript HTML
        transcript_html = self._generate_transcript_html(data)
        
        # Generate insights HTML
        insights_html = self._generate_insights_html(data)
        
        # Get visualization HTML
        timeline_html = visualizations.get('timeline', {}).get('html', '<p>Временная шкала недоступна</p>')
        heatmap_html = visualizations.get('heatmap', {}).get('html', '<p>Тепловая карта недоступна</p>')
        
        # HTML template
        html_template = f"""
        <!DOCTYPE html>
        <html lang="ru">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>ДОПРОС MVP - Комплексный анализ</title>
            <script src="https://cdn.plot.ly/plotly-latest.min.js"></script>
            <style>{custom_css}</style>
        </head>
        <body>
            <div class="header">
                <h1>🔍 ДОПРОС MVP 2.0</h1>
                <h2>Комплексный анализ допроса</h2>
                <p>Отчет сгенерирован: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}</p>
                <p>ID отчета: {report_id}</p>
            </div>
            
            <div class="section">
                <h2>📊 Краткая сводка</h2>
                <div class="stat-grid">
                    {summary_stats}
                </div>
            </div>
            
            <div class="section">
                <h2>⏱️ Временная шкала эмоций</h2>
                {timeline_html}
            </div>
            
            <div class="section">
                <h2>🌡️ Тепловая карта эмоциональной интенсивности</h2>
                {heatmap_html}
            </div>
            
            <div class="section">
                <h2>⚠️ Критические моменты</h2>
                {critical_moments_html}
            </div>
            
            <div class="section">
                <h2>📝 Транскрипция с анализом</h2>
                {transcript_html}
            </div>
            
            <div class="section">
                <h2>🧠 Психологические инсайты</h2>
                {insights_html}
            </div>
            
            <div class="section">
                <h2>🔗 Дополнительные материалы</h2>
                <ul>
                    <li><a href="{report_id}_emotions_timeline.csv">Временная шкала эмоций (CSV)</a></li>
                    <li><a href="{report_id}_transitions.csv">Анализ переходов (CSV)</a></li>
                    <li><a href="{report_id}_critical_moments.csv">Критические моменты (CSV)</a></li>
                    <li><a href="{report_id}_summary_statistics.csv">Сводная статистика (CSV)</a></li>
                </ul>
            </div>
        </body>
        </html>
        """
        
        # Write HTML file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_template)
        
        return output_path
    
    def _generate_summary_statistics(self, data: Dict[str, Any]) -> str:
        """Generate HTML for summary statistics"""
        metadata = data.get('metadata', {})
        
        # Calculate emotion distribution
        video_emotions = data.get('video_emotions', [])
        speech_emotions = data.get('speech_emotions', [])
        
        video_emotion_dist = Counter([item['emotion'] for item in video_emotions])
        speech_emotion_dist = Counter([item['emotion'] for item in speech_emotions])
        
        # Most common emotions
        most_common_video = video_emotion_dist.most_common(1)[0] if video_emotion_dist else ('нет данных', 0)
        most_common_speech = speech_emotion_dist.most_common(1)[0] if speech_emotion_dist else ('нет данных', 0)
        
        # Calculate average confidence
        video_confidences = [item['confidence'] for item in video_emotions]
        speech_confidences = [item['confidence'] for item in speech_emotions]
        
        avg_video_conf = np.mean(video_confidences) if video_confidences else 0
        avg_speech_conf = np.mean(speech_confidences) if speech_confidences else 0
        
        # Generate HTML cards
        stats_html = f"""
        <div class="stat-card">
            <div class="stat-value">{metadata.get('duration', 0):.1f}с</div>
            <div>Общая продолжительность</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{len(video_emotions)}</div>
            <div>Анализов лица</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{len(speech_emotions)}</div>
            <div>Сегментов речи</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{most_common_video[0]}</div>
            <div>Частая эмоция (лицо)</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{most_common_speech[0]}</div>
            <div>Частая эмоция (речь)</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{avg_video_conf:.2f}</div>
            <div>Средняя уверенность (лицо)</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{avg_speech_conf:.2f}</div>
            <div>Средняя уверенность (речь)</div>
        </div>
        <div class="stat-card">
            <div class="stat-value">{len(data.get('critical_moments', []))}</div>
            <div>Критических моментов</div>
        </div>
        """
        
        return stats_html
    
    def _generate_critical_moments_html(self, data: Dict[str, Any]) -> str:
        """Generate HTML for critical moments"""
        critical_moments = data.get('critical_moments', [])
        
        if not critical_moments:
            return '<p>Критические моменты не обнаружены или не анализировались.</p>'
        
        moments_html = ''
        for moment in sorted(critical_moments, key=lambda x: x.get('timestamp', 0)):
            timestamp = moment.get('timestamp', 0)
            time_str = f"{int(timestamp//60)}:{int(timestamp%60):02d}"
            importance = moment.get('importance', 5)
            description = moment.get('description', 'Описание отсутствует')
            context = moment.get('context', '')
            emotions = moment.get('emotions', [])
            
            emotion_tags = ''.join([f'<span class="emotion-tag">{emotion}</span>' for emotion in emotions])
            
            moments_html += f"""
            <div class="critical-moment">
                <h4>⚠️ {time_str} - Важность: {importance}/10</h4>
                <p><strong>Описание:</strong> {description}</p>
                {f'<p><strong>Контекст:</strong> {context}</p>' if context else ''}
                {f'<p><strong>Связанные эмоции:</strong> {emotion_tags}</p>' if emotions else ''}
            </div>
            """
        
        return moments_html
    
    def _generate_transcript_html(self, data: Dict[str, Any]) -> str:
        """Generate HTML for transcript with emotion analysis"""
        speech_emotions = data.get('speech_emotions', [])
        transcript = data.get('transcript', [])
        
        if not speech_emotions and not transcript:
            return '<p>Транскрипция недоступна.</p>'
        
        transcript_html = ''
        
        # Use speech emotions if available, otherwise fallback to transcript
        segments = speech_emotions if speech_emotions else transcript
        
        for segment in segments:
            timestamp = segment.get('timestamp', 0)
            time_str = f"{int(timestamp//60)}:{int(timestamp%60):02d}"
            text = segment.get('text', 'Текст недоступен')
            speaker = segment.get('speaker', 'Неизвестный')
            emotion = segment.get('emotion', 'неопределено')
            confidence = segment.get('confidence', 0)
            
            # Get emotion color
            emotion_color = self.emotion_translator.get_emotion_color(emotion)
            
            transcript_html += f"""
            <div class="transcript-segment">
                <div class="timestamp">{time_str} - {speaker}</div>
                <p><strong>{text}</strong></p>
                <div style="margin-top: 10px;">
                    <span class="emotion-tag" style="background-color: {emotion_color}">
                        {emotion} ({confidence:.2f})
                    </span>
                </div>
            </div>
            """
        
        return transcript_html
    
    def _generate_insights_html(self, data: Dict[str, Any]) -> str:
        """Generate HTML for psychological insights"""
        insights = data.get('openai_insights', {})
        
        if not insights:
            return '<p>Психологические инсайты недоступны. Требуется интеграция с OpenAI.</p>'
        
        insights_html = ''
        
        # General analysis
        if 'general_analysis' in insights:
            insights_html += f"""
            <h3>🎯 Общий анализ</h3>
            <p>{insights['general_analysis']}</p>
            """
        
        # Emotional patterns
        if 'emotional_patterns' in insights:
            insights_html += f"""
            <h3>📈 Эмоциональные паттерны</h3>
            <p>{insights['emotional_patterns']}</p>
            """
        
        # Behavioral indicators
        if 'behavioral_indicators' in insights:
            insights_html += f"""
            <h3>🎭 Поведенческие индикаторы</h3>
            <p>{insights['behavioral_indicators']}</p>
            """
        
        # Recommendations
        if 'recommendations' in insights:
            insights_html += f"""
            <h3>💡 Рекомендации</h3>
            <p>{insights['recommendations']}</p>
            """
        
        # Risk assessment
        if 'risk_assessment' in insights:
            insights_html += f"""
            <h3>⚡ Оценка рисков</h3>
            <p>{insights['risk_assessment']}</p>
            """
        
        return insights_html if insights_html else '<p>Инсайты обрабатываются...</p>'
    
    def generate_json_report(self, data: Dict[str, Any], report_id: str) -> str:
        """Generate comprehensive JSON report"""
        output_path = str(self.reports_dir / f"{report_id}_report.json")
        
        json_data = {
            'metadata': {
                'report_id': report_id,
                'generated_at': datetime.now().isoformat(),
                'version': 'ДОПРОС MVP 2.0',
                'duration': data.get('metadata', {}).get('duration', 0),
                'video_frames': data.get('metadata', {}).get('video_frames', 0),
                'audio_segments': data.get('metadata', {}).get('audio_segments', 0)
            },
            'video_emotions': data.get('video_emotions', []),
            'speech_emotions': data.get('speech_emotions', []),
            'critical_moments': data.get('critical_moments', []),
            'transcript': data.get('transcript', []),
            'correlations': data.get('correlations', {}),
            'statistics': data.get('statistics', {}),
            'openai_insights': data.get('openai_insights', {})
        }
        
        # Add computed statistics
        video_emotions = data.get('video_emotions', [])
        speech_emotions = data.get('speech_emotions', [])
        
        if video_emotions:
            json_data['computed_statistics'] = {
                'video_emotion_distribution': dict(Counter([item['emotion'] for item in video_emotions])),
                'video_avg_confidence': float(np.mean([item['confidence'] for item in video_emotions])),
                'video_emotion_changes': self._count_emotion_changes(video_emotions)
            }
        
        if speech_emotions:
            json_data['computed_statistics'].update({
                'speech_emotion_distribution': dict(Counter([item['emotion'] for item in speech_emotions])),
                'speech_avg_confidence': float(np.mean([item['confidence'] for item in speech_emotions])),
                'speech_emotion_changes': self._count_emotion_changes(speech_emotions)
            })
        
        # Write JSON file
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, ensure_ascii=False, indent=2)
        
        return output_path    
    def generate_pdf_report(self, data: Dict[str, Any], visualizations: Dict[str, Any], 
                          report_id: str) -> Optional[str]:
        """Generate PDF report using ReportLab"""
        
        if not REPORTLAB_AVAILABLE:
            self.logger.warning("ReportLab not available, skipping PDF generation")
            return None
        
        output_path = str(self.reports_dir / f"{report_id}_report.pdf")
        
        try:
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                textColor=colors.darkblue,
                alignment=1  # Center
            )
            
            story.append(Paragraph("ДОПРОС MVP 2.0 - Анализ допроса", title_style))
            story.append(Spacer(1, 20))
            
            # Metadata
            metadata = data.get('metadata', {})
            story.append(Paragraph(f"ID отчета: {report_id}", styles['Normal']))
            story.append(Paragraph(f"Дата генерации: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}", styles['Normal']))
            story.append(Paragraph(f"Продолжительность: {metadata.get('duration', 0):.1f} секунд", styles['Normal']))
            story.append(Spacer(1, 20))
            
            # Summary statistics
            story.append(Paragraph("Краткая статистика", styles['Heading2']))
            
            video_emotions = data.get('video_emotions', [])
            speech_emotions = data.get('speech_emotions', [])
            
            stats_data = [
                ['Параметр', 'Значение'],
                ['Кадров с эмоциями', str(len(video_emotions))],
                ['Сегментов речи', str(len(speech_emotions))],
                ['Критических моментов', str(len(data.get('critical_moments', [])))],
            ]
            
            if video_emotions:
                most_common_video = Counter([item['emotion'] for item in video_emotions]).most_common(1)[0]
                avg_video_conf = np.mean([item['confidence'] for item in video_emotions])
                stats_data.extend([
                    ['Частая эмоция (лицо)', f"{most_common_video[0]} ({most_common_video[1]} раз)"],
                    ['Средняя уверенность (лицо)', f"{avg_video_conf:.2f}"]
                ])
            
            if speech_emotions:
                most_common_speech = Counter([item['emotion'] for item in speech_emotions]).most_common(1)[0]
                avg_speech_conf = np.mean([item['confidence'] for item in speech_emotions])
                stats_data.extend([
                    ['Частая эмоция (речь)', f"{most_common_speech[0]} ({most_common_speech[1]} раз)"],
                    ['Средняя уверенность (речь)', f"{avg_speech_conf:.2f}"]
                ])
            
            table = Table(stats_data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 14),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(table)
            story.append(Spacer(1, 20))
            
            # Critical moments
            story.append(Paragraph("Критические моменты", styles['Heading2']))
            critical_moments = data.get('critical_moments', [])
            
            if critical_moments:
                for moment in sorted(critical_moments, key=lambda x: x.get('timestamp', 0)):
                    timestamp = moment.get('timestamp', 0)
                    time_str = f"{int(timestamp//60)}:{int(timestamp%60):02d}"
                    importance = moment.get('importance', 5)
                    description = moment.get('description', 'Описание отсутствует')
                    
                    story.append(Paragraph(
                        f"<b>{time_str} (важность: {importance}/10)</b>: {description}",
                        styles['Normal']
                    ))
                    story.append(Spacer(1, 10))
            else:
                story.append(Paragraph("Критические моменты не обнаружены.", styles['Normal']))
            
            story.append(Spacer(1, 20))
            
            # Insights
            insights = data.get('openai_insights', {})
            if insights:
                story.append(Paragraph("Психологические инсайты", styles['Heading2']))
                
                for key, value in insights.items():
                    if value:
                        story.append(Paragraph(f"<b>{key.replace('_', ' ').title()}:</b> {value}", styles['Normal']))
                        story.append(Spacer(1, 10))
            
            # Build PDF
            doc.build(story)
            
            return output_path
            
        except Exception as e:
            self.logger.error(f"PDF generation failed: {e}")
            return None
    
    def generate_docx_report(self, data: Dict[str, Any], report_id: str) -> Optional[str]:
        """Generate DOCX report using python-docx"""
        
        if not DOCX_AVAILABLE:
            self.logger.warning("python-docx not available, skipping DOCX generation")
            return None
        
        output_path = str(self.reports_dir / f"{report_id}_report.docx")
        
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading('ДОПРОС MVP 2.0 - Анализ допроса', 0)
            title.alignment = 1  # Center
            
            # Metadata
            doc.add_heading('Информация об отчете', level=1)
            metadata_para = doc.add_paragraph()
            metadata_para.add_run(f"ID отчета: {report_id}").bold = True
            doc.add_paragraph(f"Дата генерации: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}")
            
            metadata = data.get('metadata', {})
            doc.add_paragraph(f"Продолжительность анализа: {metadata.get('duration', 0):.1f} секунд")
            
            # Summary
            doc.add_heading('Краткая статистика', level=1)
            
            video_emotions = data.get('video_emotions', [])
            speech_emotions = data.get('speech_emotions', [])
            
            stats_table = doc.add_table(rows=1, cols=2)
            stats_table.style = 'Light Shading Accent 1'
            hdr_cells = stats_table.rows[0].cells
            hdr_cells[0].text = 'Параметр'
            hdr_cells[1].text = 'Значение'
            
            # Add statistics rows
            stats_data = [
                ('Кадров с анализом эмоций', str(len(video_emotions))),
                ('Сегментов речи', str(len(speech_emotions))),
                ('Критических моментов', str(len(data.get('critical_moments', [])))),
            ]
            
            if video_emotions:
                most_common_video = Counter([item['emotion'] for item in video_emotions]).most_common(1)[0]
                avg_video_conf = np.mean([item['confidence'] for item in video_emotions])
                stats_data.extend([
                    ('Наиболее частая эмоция (лицо)', f"{most_common_video[0]} ({most_common_video[1]} раз)"),
                    ('Средняя уверенность (лицо)', f"{avg_video_conf:.2f}")
                ])
            
            if speech_emotions:
                most_common_speech = Counter([item['emotion'] for item in speech_emotions]).most_common(1)[0]
                avg_speech_conf = np.mean([item['confidence'] for item in speech_emotions])
                stats_data.extend([
                    ('Наиболее частая эмоция (речь)', f"{most_common_speech[0]} ({most_common_speech[1]} раз)"),
                    ('Средняя уверенность (речь)', f"{avg_speech_conf:.2f}")
                ])
            
            for param, value in stats_data:
                row_cells = stats_table.add_row().cells
                row_cells[0].text = param
                row_cells[1].text = value
            
            # Critical moments
            doc.add_heading('Критические моменты', level=1)
            critical_moments = data.get('critical_moments', [])
            
            if critical_moments:
                for i, moment in enumerate(sorted(critical_moments, key=lambda x: x.get('timestamp', 0)), 1):
                    timestamp = moment.get('timestamp', 0)
                    time_str = f"{int(timestamp//60)}:{int(timestamp%60):02d}"
                    importance = moment.get('importance', 5)
                    description = moment.get('description', 'Описание отсутствует')
                    context = moment.get('context', '')
                    
                    moment_para = doc.add_paragraph()
                    moment_para.add_run(f"{i}. {time_str} (важность: {importance}/10): ").bold = True
                    moment_para.add_run(description)
                    
                    if context:
                        context_para = doc.add_paragraph()
                        context_para.add_run("Контекст: ").bold = True
                        context_para.add_run(context)
                        context_para.style = 'Intense Quote'
            else:
                doc.add_paragraph("Критические моменты не обнаружены.")
            
            # Save document
            doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            self.logger.error(f"DOCX generation failed: {e}")
            return None
    
    def _count_emotion_changes(self, emotions_list: List[Dict[str, Any]]) -> int:
        """Count emotion transitions in a list"""
        if len(emotions_list) < 2:
            return 0
        
        changes = 0
        for i in range(1, len(emotions_list)):
            if emotions_list[i]['emotion'] != emotions_list[i-1]['emotion']:
                changes += 1
        
        return changes
    
    def _calculate_correlations(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Calculate correlations between different emotion metrics"""
        correlations = {}
        
        try:
            video_emotions = data.get('video_emotions', [])
            speech_emotions = data.get('speech_emotions', [])
            
            if not video_emotions or not speech_emotions:
                return correlations
            
            # Align timestamps for correlation analysis
            video_timestamps = [item['timestamp'] for item in video_emotions]
            speech_timestamps = [item['timestamp'] for item in speech_emotions]
            
            # Find overlapping time periods
            min_time = max(min(video_timestamps), min(speech_timestamps))
            max_time = min(max(video_timestamps), max(speech_timestamps))
            
            if min_time >= max_time:
                return correlations
            
            # Create time bins for correlation analysis
            time_step = 5.0  # 5-second bins
            time_bins = np.arange(min_time, max_time + time_step, time_step)
            
            # Aggregate emotions by time bins
            video_emotion_bins = self._aggregate_emotions_by_time(video_emotions, time_bins)
            speech_emotion_bins = self._aggregate_emotions_by_time(speech_emotions, time_bins)
            
            # Calculate correlations between different metrics
            if len(video_emotion_bins) > 2 and len(speech_emotion_bins) > 2:
                # Confidence correlation
                video_conf = [bin_data.get('avg_confidence', 0) for bin_data in video_emotion_bins]
                speech_conf = [bin_data.get('avg_confidence', 0) for bin_data in speech_emotion_bins]
                
                if len(video_conf) == len(speech_conf) and len(video_conf) > 2:
                    corr_coef, p_value = pearsonr(video_conf, speech_conf)
                    correlations['video_speech_confidence'] = {
                        'correlation': float(corr_coef),
                        'p_value': float(p_value),
                        'sample_size': len(video_conf)
                    }
                
                # Emotional intensity correlation
                video_intensity = [bin_data.get('emotion_intensity', 0) for bin_data in video_emotion_bins]
                speech_intensity = [bin_data.get('emotion_intensity', 0) for bin_data in speech_emotion_bins]
                
                if len(video_intensity) == len(speech_intensity) and len(video_intensity) > 2:
                    corr_coef, p_value = pearsonr(video_intensity, speech_intensity)
                    correlations['video_speech_intensity'] = {
                        'correlation': float(corr_coef),
                        'p_value': float(p_value),
                        'sample_size': len(video_intensity)
                    }
                    
        except Exception as e:
            self.logger.warning(f"Correlation calculation failed: {e}")
        
        return correlations
    
    def _calculate_basic_statistics(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Calculate basic statistics for the data"""
        stats = {}
        
        try:
            # Video emotion statistics
            video_emotions = data.get('video_emotions', [])
            if video_emotions:
                video_emotion_names = [item['emotion'] for item in video_emotions]
                video_confidences = [item['confidence'] for item in video_emotions]
                
                stats['video_emotion_distribution'] = dict(Counter(video_emotion_names))
                stats['video_avg_confidence'] = float(np.mean(video_confidences))
                stats['video_confidence_std'] = float(np.std(video_confidences))
                stats['video_emotion_changes'] = self._count_emotion_changes(video_emotions)
            
            # Speech emotion statistics
            speech_emotions = data.get('speech_emotions', [])
            if speech_emotions:
                speech_emotion_names = [item['emotion'] for item in speech_emotions]
                speech_confidences = [item['confidence'] for item in speech_emotions]
                
                stats['speech_emotion_distribution'] = dict(Counter(speech_emotion_names))
                stats['speech_avg_confidence'] = float(np.mean(speech_confidences))
                stats['speech_confidence_std'] = float(np.std(speech_confidences))
                stats['speech_emotion_changes'] = self._count_emotion_changes(speech_emotions)
                
                # Text statistics
                texts = [item.get('text', '') for item in speech_emotions]
                stats['avg_text_length'] = float(np.mean([len(text) for text in texts]))
                stats['total_words'] = sum([len(text.split()) for text in texts])
            
            # Critical moments statistics
            critical_moments = data.get('critical_moments', [])
            if critical_moments:
                importance_scores = [moment.get('importance', 5) for moment in critical_moments]
                stats['critical_moments_count'] = len(critical_moments)
                stats['avg_importance'] = float(np.mean(importance_scores))
                stats['max_importance'] = float(max(importance_scores))
                
        except Exception as e:
            self.logger.warning(f"Statistics calculation failed: {e}")
        
        return stats
    
    def _aggregate_emotions_by_time(self, emotions_list: List[Dict[str, Any]], 
                                  time_bins: np.ndarray) -> List[Dict[str, Any]]:
        """Aggregate emotions by time bins for correlation analysis"""
        aggregated = []
        
        for i in range(len(time_bins) - 1):
            bin_start = time_bins[i]
            bin_end = time_bins[i + 1]
            
            # Find emotions in this time bin
            bin_emotions = [
                item for item in emotions_list
                if bin_start <= item['timestamp'] < bin_end
            ]
            
            if bin_emotions:
                # Calculate aggregated metrics
                confidences = [item['confidence'] for item in bin_emotions]
                emotions = [item['emotion'] for item in bin_emotions]
                
                # Simple emotion intensity mapping
                emotion_intensity_map = {
                    'злость': 0.9, 'гнев': 0.9,
                    'страх': 0.8, 
                    'отвращение': 0.7,
                    'грусть': 0.6,
                    'удивление': 0.5,
                    'радость': 0.4, 'счастье': 0.4,
                    'нейтральность': 0.1, 'спокойствие': 0.1
                }
                
                intensities = [emotion_intensity_map.get(emotion, 0.5) for emotion in emotions]
                
                bin_data = {
                    'time_start': bin_start,
                    'time_end': bin_end,
                    'emotion_count': len(bin_emotions),
                    'avg_confidence': np.mean(confidences),
                    'emotion_intensity': np.mean(intensities),
                    'dominant_emotion': Counter(emotions).most_common(1)[0][0] if emotions else 'нейтральность'
                }
            else:
                # Empty bin
                bin_data = {
                    'time_start': bin_start,
                    'time_end': bin_end,
                    'emotion_count': 0,
                    'avg_confidence': 0.0,
                    'emotion_intensity': 0.1,  # Neutral intensity
                    'dominant_emotion': 'нейтральность'
                }
            
            aggregated.append(bin_data)
        
        return aggregated
    
    def _generate_comprehensive_statistics(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """Generate comprehensive statistics for the report"""
        stats = {}
        
        try:
            # Basic counts
            video_emotions = data.get('video_emotions', [])
            speech_emotions = data.get('speech_emotions', [])
            critical_moments = data.get('critical_moments', [])
            
            stats['counts'] = {
                'video_analyses': len(video_emotions),
                'speech_segments': len(speech_emotions),
                'critical_moments': len(critical_moments),
                'total_duration': data.get('metadata', {}).get('duration', 0)
            }
            
            # Emotion distributions
            if video_emotions:
                stats['video_emotions'] = {
                    'distribution': dict(Counter([item['emotion'] for item in video_emotions])),
                    'avg_confidence': float(np.mean([item['confidence'] for item in video_emotions])),
                    'confidence_std': float(np.std([item['confidence'] for item in video_emotions])),
                    'transitions': self._count_emotion_changes(video_emotions)
                }
            
            if speech_emotions:
                stats['speech_emotions'] = {
                    'distribution': dict(Counter([item['emotion'] for item in speech_emotions])),
                    'avg_confidence': float(np.mean([item['confidence'] for item in speech_emotions])),
                    'confidence_std': float(np.std([item['confidence'] for item in speech_emotions])),
                    'transitions': self._count_emotion_changes(speech_emotions)
                }
            
            # Critical moments analysis
            if critical_moments:
                importance_scores = [moment.get('importance', 5) for moment in critical_moments]
                stats['critical_analysis'] = {
                    'avg_importance': float(np.mean(importance_scores)),
                    'max_importance': float(max(importance_scores)),
                    'min_importance': float(min(importance_scores)),
                    'importance_std': float(np.std(importance_scores))
                }
            
            # Correlations
            stats['correlations'] = data.get('correlations', {})
            
            # Time-based analysis
            if video_emotions or speech_emotions:
                all_timestamps = []
                if video_emotions:
                    all_timestamps.extend([item['timestamp'] for item in video_emotions])
                if speech_emotions:
                    all_timestamps.extend([item['timestamp'] for item in speech_emotions])
                
                if all_timestamps:
                    stats['temporal_analysis'] = {
                        'analysis_start': float(min(all_timestamps)),
                        'analysis_end': float(max(all_timestamps)),
                        'analysis_duration': float(max(all_timestamps) - min(all_timestamps)),
                        'data_points_per_minute': len(all_timestamps) / ((max(all_timestamps) - min(all_timestamps)) / 60) if max(all_timestamps) > min(all_timestamps) else 0
                    }
                    
        except Exception as e:
            self.logger.error(f"Comprehensive statistics generation failed: {e}")
        
        return stats
